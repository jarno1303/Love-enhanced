# app.py

# ============================================================================
# YMPÄRISTÖMUUTTUJIEN LATAUS - TÄYTYY OLLA ENSIMMÄISENÄ!
# ============================================================================
from dotenv import load_dotenv
load_dotenv()

# ============================================================================
# STANDARDIKIRJASTO-IMPORTIT
# ============================================================================
import os
import sqlite3
import random
import json
import re
import smtplib
import logging
import string
import secrets # Käytä secrets-moduulia salasanoihin
from dataclasses import asdict
from datetime import datetime, timedelta, timezone
from functools import wraps
from io import BytesIO
from logging.handlers import RotatingFileHandler
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

# ============================================================================
# THIRD-PARTY KIRJASTOT
# ============================================================================
from flask import Flask, jsonify, render_template, request, redirect, url_for, flash, session, abort
from flask_bcrypt import Bcrypt
# HUOM: Varmista, että User-malli models.py:ssä on päivitetty hyväksymään organization_id
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from flask_wtf.csrf import CSRFProtect, generate_csrf
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from werkzeug.middleware.proxy_fix import ProxyFix
from itsdangerous import URLSafeTimedSerializer, SignatureExpired, BadSignature

# ReportLab (PDF-generointi) - Jos käytät näitä, varmista importit
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.pdfgen import canvas
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    print("VAROITUS: ReportLab ei asennettu. PDF-vienti ei toimi.")


# python-docx (Word-dokumentit) - Jos käytät näitä, varmista importit
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("VAROITUS: python-docx ei asennettu. Word-vienti ei toimi.")


# ============================================================================
# OMAT MODUULIT
# ============================================================================
# Käytä nyt päivitettyä DatabaseManageria
# Varmista, että polku on oikein (esim. from data_access.database_manager import DatabaseManager)
try:
    from data_access.database_manager import DatabaseManager
except ImportError:
    print("KRIITTINEN VIRHE: data_access.database_manager moduulia ei löydy.")
    # Voit pysäyttää sovelluksen tässä
    import sys
    sys.exit(1)

from logic.stats_manager import EnhancedStatsManager
from logic.achievement_manager import EnhancedAchievementManager, ENHANCED_ACHIEVEMENTS
from logic.spaced_repetition import SpacedRepetitionManager
# Huom: simulation_manageria ei näytetä käytettävän suoraan, ehkä API-reiteissä?
# from logic import simulation_manager

# Varmista, että User-luokka models.py:ssä on päivitetty hyväksymään organization_id
try:
    from models.models import User, Question
except ImportError:
    print("KRIITTINEN VIRHE: models.models moduulia tai luokkia ei löydy.")
    import sys
    sys.exit(1)

from constants import DISTRACTORS

# ============================================================================
# FLASK-SOVELLUKSEN ALUSTUS
# ============================================================================

app = Flask(__name__)

# --- Konfiguraatio ---
# Hae SECRET_KEY ympäristömuuttujasta (PAKOLLINEN tuotannossa!)
SECRET_KEY = os.environ.get('SECRET_KEY')
if not SECRET_KEY:
    # Generoi väliaikainen avain kehitykseen, jos sitä ei ole asetettu
    # ÄLÄ KÄYTÄ tätä tuotannossa! Aseta aina vahva avain .env-tiedostoon.
    SECRET_KEY = secrets.token_hex(24)
    if 'pytest' not in sys.modules: # Älä tulosta testien aikana
        print("VAROITUS: SECRET_KEY ympäristömuuttuja puuttuu!")
        print("Käytetään väliaikaista avainta - ÄLÄ käytä tuotannossa!")
app.config['SECRET_KEY'] = SECRET_KEY

# Debug-tila: päällä vain jos FLASK_ENV=development
DEBUG_MODE = os.environ.get('FLASK_ENV', 'production').lower() == 'development'
app.config['DEBUG'] = DEBUG_MODE

# ProxyFix: Korjaa X-Forwarded-* headerit (Railway, Heroku, Nginx ym.)
app.config['USE_PROXY_FIX'] = True # Aktivoi ProxyFix
app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=1, x_port=1, x_prefix=1)

# Session-asetukset (turvallisuus tuotannossa)
app.config['SESSION_COOKIE_SECURE'] = not DEBUG_MODE # True tuotannossa (HTTPS)
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax' # Tai 'Strict'

# CSRF-suojaus
csrf = CSRFProtect(app)

# ============================================================================
# LOKITUS
# ============================================================================
log_level = logging.DEBUG if DEBUG_MODE else logging.INFO
log_format = '%(asctime)s %(levelname)s: %(message)s [in %(pathname)s:%(lineno)d]'

# Varmista lokihakemiston olemassaolo
log_dir = 'logs'
if not os.path.exists(log_dir):
    try:
        os.makedirs(log_dir)
    except OSError as e:
        print(f"Virhe lokihakemiston luonnissa '{log_dir}': {e}")
        # Voit päättää jatketaanko vai pysäytetäänkö

# Tiedostolokitus
log_file = os.path.join(log_dir, 'love_enhanced.log')
try:
    file_handler = RotatingFileHandler(log_file, maxBytes=10*1024*1024, backupCount=5, encoding='utf-8')
    file_handler.setFormatter(logging.Formatter(log_format))
    file_handler.setLevel(log_level)
    app.logger.addHandler(file_handler)
except Exception as e:
    print(f"Virhe tiedostolokituksen (RotatingFileHandler) alustuksessa: {e}")

# Konsolilokitus (hyödyllinen esim. Dockerissa)
stream_handler = logging.StreamHandler()
stream_handler.setFormatter(logging.Formatter(log_format))
stream_handler.setLevel(log_level)
app.logger.addHandler(stream_handler)

# Poista Flaskin oletus handler, jotta vältetään tuplalokitus
app.logger.removeHandler(app.logger.handlers[0])
app.logger.setLevel(log_level)

# Aseta myös werkzeug-lokituksen taso
logging.getLogger('werkzeug').setLevel(logging.WARNING if not DEBUG_MODE else logging.INFO)

app.logger.info(f"--- LOVe Enhanced käynnistyy (Debug Mode: {DEBUG_MODE}) ---")

# ============================================================================
# RATE LIMITING
# ============================================================================
# Käytä Redis-URL:ää tuotannossa parempaan skaalautuvuuteen
storage_uri = os.environ.get("REDIS_URL") or "memory://"
limiter = Limiter(
    get_remote_address,
    app=app,
    default_limits=["200 per hour", "50 per minute"], # Säädä tarpeen mukaan
    storage_uri=storage_uri,
    strategy="fixed-window" # Tai "moving-window"
)
app.logger.info(f"Rate Limiter alustettu (Storage: {'Redis' if 'redis' in storage_uri else 'Memory'})")

# ============================================================================
# TIETOKANTA JA MANAGERIT
# ============================================================================
try:
    # Varmista, että DatabaseManager käyttää oikeaa polkua/URL:ää
    db_manager = DatabaseManager() # Olettaa DATABASE_URL tai oletus SQLite-polku
    stats_manager = EnhancedStatsManager(db_manager)
    achievement_manager = EnhancedAchievementManager(db_manager)
    spaced_repetition_manager = SpacedRepetitionManager(db_manager)
    bcrypt = Bcrypt(app)
    app.logger.info("DatabaseManager ja muut managerit alustettu onnistuneesti.")
except Exception as e:
    app.logger.critical(f"KRIITTINEN VIRHE: Managerien alustus epäonnistui: {e}", exc_info=True)
    # Voit pysäyttää sovelluksen tässä, koska se ei voi toimia ilman tietokantaa
    import sys
    sys.exit(f"Sovelluksen käynnistys epäonnistui: Tietokantayhteysvirhe. Tarkista lokit. Virhe: {e}")

# ============================================================================
# FLASK-LOGIN SETUP (PÄIVITETTY Multi-Tenant varten)
# ============================================================================
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login_route' # Minne ohjataan, jos vaaditaan kirjautuminen
login_manager.login_message = "Ole hyvä ja kirjaudu sisään jatkaaksesi."
login_manager.login_message_category = "info"
# Session protection: 'strong' invalidates session if IP/User-Agent changes
login_manager.session_protection = "strong"

@login_manager.user_loader
def load_user(user_id):
    """
    Lataa käyttäjän tiedot ID:n perusteella session palauttamiseksi.
    Varmistaa, että käyttäjä on aktiivinen.
    """
    if user_id is None:
        return None
    try:
        user_data = db_manager.get_user_by_id(int(user_id)) # Varmista int-muunnos

        if user_data:
            # Tärkeä tarkistus: Onko käyttäjä aktiivinen?
            if user_data.get('status') != 'active':
                app.logger.warning(f"user_loader: Yritettiin ladata ei-aktiivista käyttäjää: ID {user_id}")
                return None # Älä palauta ei-aktiivista käyttäjää

            # Luo User-olio. Varmista, että User-luokka models.py:ssä tukee kaikkia näitä.
            # Lisää myös subscription_tier, jos/kun se on käytössä.
            return User(
                id=user_data['id'],
                username=user_data['username'],
                email=user_data['email'],
                role=user_data['role'], # 'user', 'admin', 'superuser'
                organization_id=user_data.get('organization_id'), # Voi olla None
                # Lisää muut tarvittavat kentät User-olioon, jos haluat ne helposti saataville
                # esim. distractors_enabled, expires_at
            )
        else:
             # Käyttäjää ei löytynyt ID:llä (saattaa tapahtua, jos käyttäjä poistettu)
             app.logger.warning(f"user_loader: Käyttäjää ei löytynyt ID:llä {user_id}")
             return None

    except ValueError:
        app.logger.warning(f"user_loader: Virheellinen user_id: {user_id}")
        return None
    except Exception as e:
        app.logger.error(f"user_loader: Odottamaton virhe käyttäjän {user_id} lataamisessa: {e}", exc_info=True)
        return None # Palauta None virhetilanteessa

@login_manager.unauthorized_handler
def unauthorized():
    """Käsittelee tilanteen, jossa vaaditaan kirjautumista."""
    # Jos pyyntö oli API:lle, palauta JSON-virhe
    if request.path.startswith('/api/'):
        return jsonify(error="Kirjautuminen vaaditaan"), 401
    # Muuten ohjaa kirjautumissivulle ja tallenna alkuperäinen kohde
    flash("Ole hyvä ja kirjaudu sisään jatkaaksesi.", "info")
    return redirect(url_for('login_route', next=request.url))


# ============================================================================
# APUFUNKTIOT (Dekoraattorit ja muut)
# ============================================================================

def admin_required(f):
    """
    Dekoraattori, joka vaatii vähintään 'admin'-roolin (sallii myös 'superuser').
    Varmistaa myös, että 'admin'-roolilla on organisaatio.
    """
    @wraps(f)
    @login_required # Varmista ensin, että käyttäjä on kirjautunut
    def decorated_function(*args, **kwargs):
        # Tarkista rooli
        if not hasattr(current_user, 'role') or current_user.role not in ['admin', 'superuser']:
            app.logger.warning(f"Pääsy estetty (admin/superuser vaadittu): Käyttäjä {getattr(current_user, 'username', 'N/A')} (Rooli: {getattr(current_user, 'role', 'Ei roolia')}) yritti polkuun {request.path}")
            flash("Pääsy kielletty. Vaatii ylläpitäjän oikeudet.", "danger")
            # Ohjaa käyttäjän omaan dashboardiin
            return redirect(url_for('dashboard_route'))

        # Lisätarkistus: Varmista, että 'admin'-roolilla on organisaatio
        if current_user.role == 'admin' and (not hasattr(current_user, 'organization_id') or current_user.organization_id is None):
             app.logger.error(f"Admin-käyttäjä {current_user.username} (ID: {current_user.id}) ilman organisaatiota yritti polkuun {request.path}. Estetty.")
             flash("Virheellinen konfiguraatio: Sinua ei ole liitetty organisaatioon. Ota yhteys pääkäyttäjään.", "danger")
             # Kirjaa ulos turvallisuussyistä, koska tili on väärin konfiguroitu
             logout_user()
             return redirect(url_for('login_route'))

        # Jos tarkistukset ok, suorita alkuperäinen funktio
        return f(*args, **kwargs)
    return decorated_function

def superuser_required(f):
    """
    Dekoraattori, joka vaatii 'superuser'-roolin.
    """
    @wraps(f)
    @login_required # Varmista ensin kirjautuminen
    def decorated_function(*args, **kwargs):
        # Tarkista rooli
        if not hasattr(current_user, 'role') or current_user.role != 'superuser':
            app.logger.warning(f"Pääsy estetty (superuser vaadittu): Käyttäjä {getattr(current_user, 'username', 'N/A')} (Rooli: {getattr(current_user, 'role', 'Ei roolia')}) yritti polkuun {request.path}")
            flash("Pääsy kielletty. Vaatii järjestelmän pääkäyttäjän oikeudet.", "danger")
            # Ohjaa admin-dashboardiin (joka sitten voi ohjata superuser-näkymään, jos rooli olisi oikea)
            # Tai ohjaa suoraan käyttäjän omaan dashboardiin
            if hasattr(current_user, 'role') and current_user.role == 'admin':
                return redirect(url_for('admin_route'))
            else:
                return redirect(url_for('dashboard_route'))

        # Jos rooli on oikea, suorita funktio
        return f(*args, **kwargs)
    return decorated_function

# --- Muut aiemmat apufunktiot (generate_secure_password, salasanan palautus jne.) ---
# (Nämä pysyvät samoina kuin edellisessä vastauksessa)
def generate_secure_password(length=12):
    """Luo turvallisen satunnaisen salasanan (kirjaimia, numeroita)."""
    if length < 10: length = 10
    chars = string.ascii_letters + string.digits
    while True:
        password = ''.join(secrets.choice(chars) for _ in range(length))
        if (any(c.islower() for c in password)
                and any(c.isupper() for c in password)
                and any(c.isdigit() for c in password)):
            return password

def generate_reset_token(email, expires_sec=1800):
    """Generoi aikarajoitetun tokenin salasanan palautukseen."""
    serializer = URLSafeTimedSerializer(app.config['SECRET_KEY'])
    return serializer.dumps(email, salt='password-reset-salt')

def verify_reset_token(token, max_age_sec=1800):
    """Varmistaa salasanan palautustokenin."""
    serializer = URLSafeTimedSerializer(app.config['SECRET_KEY'])
    try:
        email = serializer.loads(token, salt='password-reset-salt', max_age=max_age_sec)
        return email
    except (SignatureExpired, BadSignature):
        return None

def send_reset_email(user_email, reset_url):
    """Lähettää salasanan palautusviestin."""
    BREVO_API_KEY = os.environ.get('BREVO_API_KEY')
    FROM_EMAIL = os.environ.get('FROM_EMAIL', 'noreply@loveenhanced.fi')

    if not BREVO_API_KEY or DEBUG_MODE:
        app.logger.warning(f"Sähköpostipalvelua (Brevo) ei konfiguroitu tai DEBUG. Palautuslinkki ({user_email}): {reset_url}")
        if DEBUG_MODE: return True
        return False

    import requests
    url = "https://api.brevo.com/v3/smtp/email"
    headers = {"accept": "application/json", "api-key": BREVO_API_KEY, "content-type": "application/json"}
    # Tässä voisi käyttää templatea email_html:n luomiseen
    email_html = render_template('emails/reset_password_email.html', reset_url=reset_url) # Olettaa templaten olemassaolon

    payload = {
        "sender": {"name": "LOVe Enhanced", "email": FROM_EMAIL},
        "to": [{"email": user_email}],
        "subject": "LOVe Enhanced - Salasanan palautuspyyntö",
        "htmlContent": email_html
    }
    try:
        response = requests.post(url, json=payload, headers=headers)
        response.raise_for_status()
        app.logger.info(f"Salasanan palautusviesti lähetetty: {user_email}")
        return True
    except requests.exceptions.RequestException as e:
        app.logger.error(f"Sähköpostin lähetys epäonnistui (Brevo): {e}", exc_info=True)
        if response is not None:
             app.logger.error(f"Brevo response: {response.status_code}, {response.text}")
        return False

# ============================================================================
# --- API-REITIT ---
# ============================================================================

# Vastaa GET /api/csrf-token pyyntöön CSRF-tokenilla
@app.route('/api/csrf-token')
def get_csrf_token():
    """Palauttaa CSRF-tokenin käytettäväksi AJAX-pyynnöissä."""
    return jsonify({'csrf_token': generate_csrf()})


# --- Kysymykset ---
@app.route("/api/questions")
@login_required # Vaatii kirjautumisen
@limiter.limit("60 per minute") # Rajoita kyselytiheyttä
def get_questions_api():
    """
    Hakee harjoituskysymyksiä valintojen mukaan.
    Tukee kategorioita, vaikeustasoja ja määrää.
    Simulaatiotila hakee 50 kysymystä ilman suodatusta.
    """
    try:
        # Hae parametrit turvallisesti oletusarvoilla
        categories = request.args.getlist('categories') or None
        difficulties = request.args.getlist('difficulties') or None
        limit = request.args.get('count', default=10, type=int)
        simulation = request.args.get('simulation') == 'true'

        # Rajoita maksimimäärää järkeväksi (estää DoS)
        limit = max(1, min(limit, 100)) # 1-100 kysymystä kerralla

        app.logger.debug(f"API /api/questions: user={current_user.id}, sim={simulation}, cat={categories}, diff={difficulties}, limit={limit}")

        # Freemium-rajoitus (lisää myöhemmin)
        # if current_user.subscription_tier == 'free': ...

        questions = []
        if simulation:
            # Simulaatiotila: Hae 50 satunnaista validoitua ID:tä
            question_ids = db_manager.get_random_question_ids(limit=50)
            if not question_ids:
                 app.logger.warning(f"Ei kysymyksiä simulaatioon käyttäjälle {current_user.id}")
                 return jsonify({'questions': [], 'message': 'Ei riittävästi kysymyksiä simulaatioon.'}), 200

            # Hae kysymykset ID-listan perusteella käyttäjän progressin kanssa
            questions_objs = [db_manager.get_question_by_id(qid, current_user.id) for qid in question_ids]
            questions = [q for q in questions_objs if q is not None] # Suodata pois mahdolliset None-arvot
            app.logger.info(f"Haettiin {len(questions)} kysymystä simulaatioon käyttäjälle {current_user.id}")
        else:
            # Normaali harjoittelu: Hae suodatettuja kysymyksiä
            questions = db_manager.get_questions(
                user_id=current_user.id,
                categories=categories,
                difficulties=difficulties,
                limit=limit
            )
            app.logger.info(f"Haettiin {len(questions)} harjoituskysymystä käyttäjälle {current_user.id}")


        # Prosessoi kysymykset JSON-muotoon
        # HUOM: ÄLÄ paljasta 'correct'-indeksiä tai 'explanation'-kenttää tässä vaiheessa!
        # Ne palautetaan vasta /api/submit_answer -vastauksessa.
        questions_list = []
        for q in questions:
            if not isinstance(q, Question): continue # Varmistus
            q_dict = asdict(q) # Jos Question on dataclass
            # Poista arkaluontoiset tiedot ennen lähetystä clientille
            q_dict.pop('correct', None)
            q_dict.pop('explanation', None)
            q_dict.pop('question_normalized', None) # Ei tarvita clientilla
            # Sekoita vastausvaihtoehdot vasta client-puolella, jos halutaan
            # TAI sekoita tässä ja tallenna sekoitettu järjestys sessioon/clientille
            questions_list.append(q_dict)

        return jsonify({'questions': questions_list})

    except ValueError as ve:
        app.logger.warning(f"Virheellinen parametri /api/questions: {ve}")
        return jsonify({'error': 'Virheellinen parametri (esim. count).', 'details': str(ve)}), 400
    except Exception as e:
        app.logger.error(f"Virhe /api/questions haussa: {e}", exc_info=True)
        return jsonify({'error': 'Palvelinvirhe kysymyksiä haettaessa.'}), 500


@app.route("/api/review-questions")
@login_required
@limiter.limit("30 per minute")
def get_review_questions_api():
    """Hakee seuraavan erääntyneen kertauskysymyksen ja mahdollisesti häiriötekijän."""
    try:
        # Hae seuraava erääntynyt kysymys Spaced Repetition Managerilta
        # limit=1 hakee vain yhden kerrallaan
        due_questions = spaced_repetition_manager.get_due_questions(current_user.id, limit=1)

        question_data = None
        if due_questions:
            question_obj = due_questions[0]
            if isinstance(question_obj, Question):
                question_data = asdict(question_obj)
                # Poista arkaluontoiset tiedot tästäkin
                question_data.pop('correct', None)
                question_data.pop('explanation', None)
                question_data.pop('question_normalized', None)
            else:
                 app.logger.error(f"get_due_questions palautti virheellisen objektin: {type(question_obj)}")


        distractor_data = None
        # Lisää häiriötekijä todennäköisyyden perusteella
        # Varmista, että current_user-objektissa on nämä tiedot
        if hasattr(current_user, 'distractors_enabled') and current_user.distractors_enabled and \
           hasattr(current_user, 'distractor_probability'):
            probability = current_user.distractor_probability / 100.0
            if random.random() < probability:
                distractor_data = random.choice(DISTRACTORS)
                app.logger.info(f"Näytetään häiriötekijä kertauksessa käyttäjälle {current_user.id} (P={probability*100:.1f}%)")

        return jsonify({'question': question_data, 'distractor': distractor_data})

    except Exception as e:
        app.logger.error(f"Virhe /api/review-questions haussa: {e}", exc_info=True)
        return jsonify({'question': None, 'distractor': None, 'error': 'Palvelinvirhe'}), 500


# --- Vastaukset ---
@app.route("/api/submit_answer", methods=['POST'])
@login_required
@limiter.limit("100 per minute") # Salli tiheämpi vastaaminen
def submit_answer_api():
    """Vastaanottaa vastauksen, tarkistaa sen, päivittää tilastot ja SR-tiedot."""
    data = request.get_json()
    if not data:
        return jsonify({'error': 'Pyynnön body puuttuu tai ei ole JSON.'}), 400

    question_id = data.get('question_id')
    # Käytä valittua tekstiä tai indeksiä tarkistukseen riippuen clientin toteutuksesta
    # Oletetaan tässä, että client lähettää valitun VAIHTOEHDON TEKSTIN
    selected_option_text = data.get('selected_option_text')
    time_taken_str = data.get('time_taken', '0') # Ota vastaan merkkijonona

    # Validoi syötteet
    try:
        question_id = int(question_id)
        # Muunna aika liukuluvuksi (sekunteina)
        time_taken = float(time_taken_str)
        if time_taken < 0: time_taken = 0 # Estä negatiiviset ajat
    except (ValueError, TypeError):
        app.logger.warning(f"Virheellinen data /api/submit_answer: q_id={question_id}, time={time_taken_str}")
        return jsonify({'error': 'Virheellinen question_id tai time_taken.'}), 400

    if selected_option_text is None:
         return jsonify({'error': 'selected_option_text puuttuu.'}), 400


    # Hae kysymyksen täydet tiedot (sis. oikean vastauksen) tietokannasta
    # Käytä user_id:tä, jotta saadaan myös SR-tiedot samalla
    question = db_manager.get_question_by_id(question_id, current_user.id)

    if not question:
        app.logger.warning(f"Vastaus tuntemattomaan kysymykseen: ID {question_id}, Käyttäjä {current_user.id}")
        return jsonify({'error': 'Kysymystä ei löytynyt.'}), 404

    # Tarkista vastaus
    try:
        correct_option_text = question.options[question.correct]
        is_correct = (selected_option_text == correct_option_text)
    except IndexError:
        app.logger.error(f"IndexError tarkistaessa vastausta kysymykseen {question_id}. Correct index: {question.correct}, Options: {question.options}")
        return jsonify({'error': 'Virhe kysymyksen datassa (väärä indeksi).'}), 500
    except Exception as check_e:
         app.logger.error(f"Odottamaton virhe vastauksen tarkistuksessa: {check_e}", exc_info=True)
         return jsonify({'error': 'Virhe vastausta tarkistaessa.'}), 500

    # --- Päivitä tietokanta ---
    try:
        # 1. Päivitä normaalit tilastot (attempts, progress)
        # Tämä hoitaa UPSERTin user_question_progressiin ja INSERTin question_attempts-tauluun
        db_manager.update_question_stats(question_id, is_correct, time_taken, current_user.id)

        # 2. Päivitä Spaced Repetition -järjestelmä
        # Määritä suorituksen laatu (0-5 asteikolla SR-algoritmille)
        # Voit hienosäätää tätä: esim. nopea oikea vastaus = 5, hidas oikea = 4, väärä = 1 tai 2
        quality = 5 if is_correct else 1 # Yksinkertainen malli: oikein=5, väärin=1

        # Laske uusi intervalli ja ease_factor käyttäen SR-manageria
        # question-objekti sisältää vanhat `interval` ja `ease_factor` arvot
        new_interval, new_ease_factor = spaced_repetition_manager.calculate_next_review(
            question=question, # Välitetään Question-objekti, jossa vanhat arvot
            performance_rating=quality
        )

        # Tallenna päivitetyt SR-tiedot tietokantaan
        spaced_repetition_manager.record_review(
            user_id=current_user.id,
            question_id=question_id,
            interval=new_interval,
            ease_factor=new_ease_factor
        )
        app.logger.debug(f"SR päivitetty: U={current_user.id}, Q={question_id}, Qlty={quality}, NInt={new_interval}, NEa={new_ease_factor:.2f}")

    except Exception as db_update_e:
        # Lokita kriittinen virhe, mutta yritä silti palauttaa vastaus clientille
        app.logger.error(f"Kriittinen virhe tietokannan päivityksessä vastauksen jälkeen: {db_update_e}", exc_info=True)
        # Älä palauta 500, koska vastaus itsessään on käsitelty. Client voi näyttää palautteen.
        # Voit lisätä vastaukseen tiedon tallennusongelmasta.
        # return jsonify({'error': 'Virhe tietojen tallennuksessa.'}), 500

    # --- Tarkista saavutukset ---
    new_achievement_ids = []
    try:
        # Voit välittää kontekstia, esim. oliko vastaus nopea
        context = {'time_taken': time_taken, 'is_correct': is_correct}
        new_achievement_ids = achievement_manager.check_achievements(current_user.id, context=context)
    except Exception as ach_e:
        app.logger.error(f"Virhe saavutusten tarkistuksessa: {ach_e}", exc_info=True)

    new_achievements_details = []
    if new_achievement_ids:
        # Hae avattujen saavutusten tiedot näytettäväksi
        for ach_id in new_achievement_ids:
            if ach_id in ENHANCED_ACHIEVEMENTS:
                 ach_obj = ENHANCED_ACHIEVEMENTS[ach_id]
                 # Muunna dataclass dictiksi tai hae tiedot manuaalisesti
                 try:
                      new_achievements_details.append(asdict(ach_obj))
                 except TypeError: # Jos ei ole dataclass
                      new_achievements_details.append({
                           'id': getattr(ach_obj, 'id', ach_id),
                           'name': getattr(ach_obj, 'name', 'N/A'),
                           'description': getattr(ach_obj, 'description', ''),
                           'icon': getattr(ach_obj, 'icon', '🏆')
                      })
        app.logger.info(f"Käyttäjä {current_user.username} avasi {len(new_achievements_details)} saavutusta: {new_achievement_ids}")


    # Palauta vastaus clientille, sisältäen oikean vastauksen ja selityksen
    return jsonify({
        'correct': is_correct,
        'correct_answer_index': question.correct, # Indeksi clientille korostusta varten
        'correct_answer_text': correct_option_text, # Teksti varmuuden vuoksi
        'explanation': question.explanation,
        'new_achievements': new_achievements_details # Lista avatuista saavutuksista
    })


@app.route("/api/submit_distractor", methods=['POST'])
@login_required
@limiter.limit("100 per minute")
def submit_distractor_api():
    """Vastaanottaa vastauksen häiriötekijään ja tallentaa sen."""
    data = request.get_json()
    if not data:
        return jsonify({'error': 'Pyynnön body puuttuu.'}), 400

    scenario = data.get('scenario')
    user_choice_str = data.get('user_choice') # Indeksi (0, 1, 2...)
    response_time_str = data.get('response_time', '0') # Millisekunteina

    # Validoi
    try:
        user_choice = int(user_choice_str)
        response_time = int(response_time_str)
        if response_time < 0: response_time = 0
    except (ValueError, TypeError):
        return jsonify({'error': 'Virheellinen user_choice tai response_time.'}), 400

    if not scenario:
        return jsonify({'error': 'Skenaario puuttuu.'}), 400

    # Etsi oikea vastaus DISTRACTORS-listasta
    correct_choice = -1 # Oletus, jos skenaariota ei löydy
    feedback_text = "Tuntematon skenaario."
    distractor_found = False
    for dist in DISTRACTORS:
        if dist['scenario'] == scenario:
            correct_choice = dist.get('correct', -1)
            feedback_text = dist.get('feedback', {}).get(user_choice, "Ei palautetta tälle valinnalle.")
            distractor_found = True
            break

    if not distractor_found:
         app.logger.warning(f"Vastaus tuntemattomaan häiriötekijäskenaarioon: '{scenario}'")
         # Voit silti tallentaa yrityksen tai palauttaa virheen
         # return jsonify({'error': 'Tuntematon skenaario.'}), 400

    is_correct = (user_choice == correct_choice)

    # Tallenna yritys tietokantaan
    try:
        success = db_manager._execute("""
            INSERT INTO distractor_attempts
            (user_id, distractor_scenario, user_choice, correct_choice, is_correct, response_time, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?)
        """, (current_user.id, scenario, user_choice, correct_choice, is_correct, response_time, datetime.now()), fetch='none')
        # _execute ei palauta arvoa, oletetaan onnistuneeksi jos ei poikkeusta
        app.logger.info(f"Käyttäjä {current_user.username} vastasi häiriötekijään '{scenario[:30]}...': {'Oikein' if is_correct else 'Väärin'}")

    except Exception as e:
        app.logger.error(f"Virhe häiriötekijäyrityksen tallennuksessa: {e}", exc_info=True)
        # Älä palauta 500 clientille, koska palaute voidaan silti näyttää
        # return jsonify({'error': 'Virhe tallennuksessa.'}), 500

    # Palauta

# --- Tilastot ja Edistyminen ---

@app.route("/api/stats")
@login_required
@limiter.limit("30 per minute") # Näitä ei tarvitse hakea jatkuvasti
def get_stats_api():
    """Hakee kattavat oppimistilastot kirjautuneelle käyttäjälle."""
    try:
        analytics = stats_manager.get_learning_analytics(current_user.id)
        return jsonify(analytics)
    except Exception as e:
        app.logger.error(f"Virhe /api/stats haussa käyttäjälle {current_user.id}: {e}", exc_info=True)
        return jsonify({'error': 'Virhe tilastojen haussa.'}), 500

@app.route("/api/distractor_stats")
@login_required
@limiter.limit("30 per minute")
def get_distractor_stats_api():
    """Palauttaa häiriötekijätilastot käyttäjälle."""
    try:
        # Hae tilastot tietokannasta
        query = """
            SELECT
                COUNT(*) as total_attempts,
                SUM(CASE WHEN is_correct THEN 1 ELSE 0 END) as correct_attempts,
                AVG(CASE WHEN response_time > 0 THEN response_time ELSE NULL END) as avg_response_time_ms
            FROM distractor_attempts
            WHERE user_id = ?
        """
        result = db_manager._execute(query, (current_user.id,), fetch='one')

        if not result or result.get('total_attempts', 0) == 0:
            stats = {
                'total_attempts': 0,
                'correct_attempts': 0,
                'success_rate': 0.0,
                'avg_response_time_ms': 0
            }
        else:
            total = result['total_attempts']
            correct = result.get('correct_attempts', 0) or 0 # Varmista 0, jos SUM palauttaa NULL
            stats = {
                'total_attempts': total,
                'correct_attempts': correct,
                'success_rate': round((correct / total * 100) if total > 0 else 0, 1),
                'avg_response_time_ms': round(result.get('avg_response_time_ms') or 0)
            }

        # Voit halutessasi hakea myös jakauman skenaarioittain tai viimeisimmät yritykset
        # recent_query = """ SELECT ... FROM distractor_attempts ... LIMIT 5"""
        # recent_attempts = db_manager._execute(recent_query, (current_user.id,), fetch='all')

        return jsonify(stats)

    except Exception as e:
        app.logger.error(f"Virhe /api/distractor_stats haussa käyttäjälle {current_user.id}: {e}", exc_info=True)
        # Palauta oletusarvot virhetilanteessa
        return jsonify({
            'total_attempts': 0, 'correct_attempts': 0, 'success_rate': 0.0, 'avg_response_time_ms': 0,
            'error': 'Virhe tilastojen haussa.'
        }), 500


@app.route("/api/incorrect_questions")
@login_required
@limiter.limit("30 per minute")
def get_incorrect_questions_api():
    """Hakee kysymykset, joihin käyttäjä on vastannut väärin, piilottaen kuitatut."""
    try:
        # Käytä boolean-arvoa PostgreSQL:lle, integer SQLite:lle
        false_val = False if db_manager.is_postgres else 0

        # Hae kysymykset, joissa virheitä JA joita ei ole kuitattu
        incorrect_questions = db_manager._execute("""
            SELECT
                q.id, q.question, q.category, q.difficulty,
                p.times_shown, p.times_correct, p.last_shown,
                -- Laske onnistumisprosentti turvallisesti (vältä nollalla jako)
                CASE
                    WHEN p.times_shown > 0 THEN ROUND((CAST(p.times_correct AS REAL) * 100.0) / p.times_shown, 1)
                    ELSE 0.0
                END as success_rate
            FROM questions q
            JOIN user_question_progress p ON q.id = p.question_id
            WHERE p.user_id = ?
              AND p.times_correct < p.times_shown -- Enemmän näyttökertoja kuin oikeita vastauksia
              AND (p.mistake_acknowledged IS NULL OR p.mistake_acknowledged = ?) -- Ei kuitattu
            ORDER BY success_rate ASC, p.last_shown DESC NULLS LAST
            LIMIT 100 -- Lisää raja, ettei palauteta liikaa kerralla
        """, (current_user.id, false_val), fetch='all')

        return jsonify({'questions': incorrect_questions if incorrect_questions else []})

    except Exception as e:
        app.logger.error(f"Virhe /api/incorrect_questions haussa käyttäjälle {current_user.id}: {e}", exc_info=True)
        return jsonify({'error': 'Virhe kehityskohteiden haussa.'}), 500

@app.route("/api/mistakes/acknowledge", methods=['POST'])
@login_required
@limiter.limit("60 per minute") # Salli useamman kuittaus
def acknowledge_mistakes_api():
    """Merkitsee yhden tai useamman kehityskohteen kuitatuksi."""
    data = request.get_json()
    question_ids = data.get('question_ids')

    if not isinstance(question_ids, list) or not question_ids:
        return jsonify({'success': False, 'error': 'question_ids-lista puuttuu tai on tyhjä.'}), 400

    # Varmista, että ID:t ovat numeroita
    try:
        # Poista mahdolliset duplikaatit ja varmista int-muoto
        question_ids_int = list(set(int(qid) for qid in question_ids))
        if not question_ids_int:
             return jsonify({'success': False, 'error': 'question_ids-lista ei sisältänyt kelvollisia numeroita.'}), 400
    except ValueError:
        return jsonify({'success': False, 'error': 'question_ids sisältää virheellisiä arvoja.'}), 400


    try:
        # Käytä oikeaa boolean-arvoa
        true_val = True if db_manager.is_postgres else 1

        # Rakenna kysely turvallisesti (älä käytä f-stringiä listoille!)
        placeholders = ','.join([db_manager.param_style] * len(question_ids_int))
        query = f"""
            UPDATE user_question_progress
            SET mistake_acknowledged = ?
            WHERE user_id = ? AND question_id IN ({placeholders})
        """
        params = [true_val, current_user.id] + question_ids_int

        db_manager._execute(query, tuple(params)) # Suorita päivitys

        # Tässä voisi tarkistaa, montako riviä päivitettiin, mutta _execute ei palauta sitä
        acknowledged_count = len(question_ids_int) # Oletetaan, että kaikki löytyivät

        app.logger.info(f"Käyttäjä {current_user.id} kuittasi {acknowledged_count} kehityskohdetta: {question_ids_int}")
        return jsonify({'success': True, 'acknowledged_count': acknowledged_count})

    except Exception as e:
        app.logger.error(f"Virhe kehityskohteiden kuittauksessa käyttäjälle {current_user.id}: {e}", exc_info=True)
        return jsonify({'success': False, 'error': 'Virhe kuittauksessa.'}), 500


@app.route("/api/question_progress/<int:question_id>")
@login_required
@limiter.limit("60 per minute")
def get_question_progress_api(question_id):
    """Hakee käyttäjän edistymisen tietyssä kysymyksessä."""
    try:
        progress = db_manager._execute("""
            SELECT
                times_shown,
                times_correct,
                last_shown,
                CASE
                    WHEN times_shown > 0 THEN ROUND((CAST(times_correct AS REAL) * 100.0) / times_shown, 1)
                    ELSE 0.0
                END as success_rate,
                mistake_acknowledged
            FROM user_question_progress
            WHERE user_id = ? AND question_id = ?
        """, (current_user.id, question_id), fetch='one')

        if progress:
            # Varmista boolean-muoto mistake_acknowledged-kentälle
            progress_dict = dict(progress)
            progress_dict['mistake_acknowledged'] = bool(progress_dict.get('mistake_acknowledged', False))
            return jsonify(progress_dict)
        else:
            # Palauta oletusarvot, jos progressia ei löydy
            return jsonify({
                'times_shown': 0,
                'times_correct': 0,
                'success_rate': 0.0,
                'last_shown': None,
                'mistake_acknowledged': False
            })

    except Exception as e:
        app.logger.error(f"Virhe kysymyksen {question_id} edistymisen haussa käyttäjälle {current_user.id}: {e}", exc_info=True)
        return jsonify({'error': 'Virhe edistymisen haussa.'}), 500


# --- Saavutukset ---
@app.route("/api/achievements")
@login_required
@limiter.limit("30 per minute")
def get_achievements_api():
    """Hakee kaikki saavutukset ja käyttäjän edistymisen niissä."""
    try:
        # Hae käyttäjän avaamat saavutukset (ID:t ja ajankohdat)
        unlocked_raw = db_manager.get_user_achievements(current_user.id) # Palauttaa listan dict-objekteja
        unlocked_map = {item['achievement_id']: item['unlocked_at'] for item in unlocked_raw} if unlocked_raw else {}

        all_achievements_list = []
        # Käy läpi kaikki mahdolliset saavutukset ENHANCED_ACHIEVEMENTS-listasta
        for ach_id, ach_details in ENHANCED_ACHIEVEMENTS.items():
            is_unlocked = ach_id in unlocked_map
            ach_data = {
                'id': ach_id,
                'name': ach_details.name,
                'description': ach_details.description,
                'icon': ach_details.icon,
                'unlocked': is_unlocked,
                'unlocked_at': unlocked_map.get(ach_id) if is_unlocked else None,
                # Lisää progress-tieto, jos AchievementManager tarjoaa sen
                # 'progress': achievement_manager.get_achievement_progress(current_user.id, ach_id)
            }
            all_achievements_list.append(ach_data)

        # Voit halutessasi järjestää listan esim. avattujen mukaan
        # all_achievements_list.sort(key=lambda x: (not x['unlocked'], x['name']))

        return jsonify(all_achievements_list)
    except Exception as e:
        app.logger.error(f"Virhe /api/achievements haussa käyttäjälle {current_user.id}: {e}", exc_info=True)
        # Palauta tyhjä lista virhetilanteessa, jotta UI ei kaadu
        return jsonify([]), 500


# --- Asetukset ---
@app.route("/api/settings/toggle_distractors", methods=['POST'])
@login_required
@limiter.limit("20 per hour") # Asetuksia ei muutella jatkuvasti
def toggle_distractors_api():
    """Ottaa häiriötekijät käyttöön tai pois käytöstä."""
    data = request.get_json()
    is_enabled = data.get('enabled', False) # Ota boolean-arvo

    try:
        # Käytä db_manager.update_user -funktiota
        success, error = db_manager.update_user(current_user.id, {'distractors_enabled': bool(is_enabled)})
        if success:
            app.logger.info(f"Käyttäjä {current_user.username} muutti häiriötekijäasetuksen: {'Päällä' if is_enabled else 'Pois'}")
            # Päivitä myös current_user-objekti sessiossa, jos mahdollista (vaatii User-luokan muokkauksen)
            # current_user.distractors_enabled = bool(is_enabled)
            return jsonify({'success': True, 'distractors_enabled': bool(is_enabled)})
        else:
             app.logger.error(f"Virhe häiriötekijäasetuksen päivityksessä käyttäjälle {current_user.id}: {error}")
             return jsonify({'success': False, 'error': error or 'Päivitys epäonnistui.'}), 500
    except Exception as e:
        app.logger.error(f"Odottamaton virhe /api/settings/toggle_distractors: {e}", exc_info=True)
        return jsonify({'success': False, 'error': 'Palvelinvirhe.'}), 500

@app.route("/api/settings/update_distractor_probability", methods=['POST'])
@login_required
@limiter.limit("20 per hour")
def update_distractor_probability_api():
    """Päivittää häiriötekijöiden näyttötodennäköisyyden."""
    data = request.get_json()
    probability_str = data.get('probability', '25') # Ota vastaan merkkijonona

    try:
        probability = int(probability_str)
        # Rajoita arvo välille 0-100
        probability = max(0, min(100, probability))

        success, error = db_manager.update_user(current_user.id, {'distractor_probability': probability})
        if success:
            app.logger.info(f"Käyttäjä {current_user.username} päivitti häiriötekijöiden todennäköisyyden: {probability}%")
            # Päivitä current_user-objekti
            # current_user.distractor_probability = probability
            return jsonify({'success': True, 'probability': probability})
        else:
             app.logger.error(f"Virhe todennäköisyysasetuksen päivityksessä käyttäjälle {current_user.id}: {error}")
             return jsonify({'success': False, 'error': error or 'Päivitys epäonnistui.'}), 500

    except ValueError:
        return jsonify({'success': False, 'error': 'Virheellinen todennäköisyysarvo.'}), 400
    except Exception as e:
        app.logger.error(f"Odottamaton virhe /api/settings/update_distractor_probability: {e}", exc_info=True)
        return jsonify({'success': False, 'error': 'Palvelinvirhe.'}), 500


@app.route('/api/user_preferences', methods=['POST'])
@login_required
@limiter.limit("10 per hour")
def save_user_preferences_api():
    """Tallentaa käyttäjän oletuskategoriat ja -vaikeustasot harjoitteluun."""
    data = request.get_json()
    categories = data.get('categories', [])
    difficulties = data.get('difficulties', [])

    # Validoi data (varmista, että ne ovat listoja)
    if not isinstance(categories, list) or not isinstance(difficulties, list):
        return jsonify({'status': 'error', 'message': 'Virheellinen dataformaatti.'}), 400

    # Tallenna tietokantaan
    success, error = db_manager.update_user_practice_preferences(current_user.id, categories, difficulties)

    if success:
        app.logger.info(f"Käyttäjä {current_user.id} tallensi harjoitteluasetukset: C={categories}, D={difficulties}")
        return jsonify({'status': 'success', 'message': 'Asetukset tallennettu.'}), 200
    else:
        app.logger.error(f"Virhe harjoitteluasetusten tallennuksessa käyttäjälle {current_user.id}: {error}")
        return jsonify({'status': 'error', 'message': error or 'Tallennus epäonnistui.'}), 500
    
# --- Simulaatio API ---
# Nämä reitit käyttävät Flaskin server-puolen sessiota simulaation tilan tallentamiseen.

@app.route('/api/simulation/start', methods=['POST'])
@login_required
# @premium_required # Lisää tämä, jos simulaatio on maksullinen
@limiter.limit("5 per hour") # Rajoita uusien simulaatioiden aloittamista
def start_simulation_api():
    """Aloittaa uuden koesimulaation ja tallentaa sen tiedot sessioon."""
    try:
        # Tarkista, onko jo aktiivinen simulaatio
        if 'simulation' in session and session['simulation'].get('user_id') == current_user.id:
            # Voit päättää, sallitaanko uuden aloitus vai pakotetaanko jatkamaan vanhaa
            app.logger.warning(f"Käyttäjä {current_user.id} yritti aloittaa uuden simulaation, vaikka vanha oli kesken.")
            # Palauta tieto olemassaolevasta sessiosta?
            # return jsonify({'error': 'Aktiivinen simulaatio on jo käynnissä.', 'resume': True}), 409

            # Tai poista vanha ja aloita uusi:
            session.pop('simulation', None)
            app.logger.info(f"Poistettiin vanha kesken ollut simulaatio käyttäjälle {current_user.id}")


        # Hae 50 satunnaista kysymys-ID:tä
        question_ids = db_manager.get_random_question_ids(limit=50)

        # Tarkista, saatiinko tarpeeksi kysymyksiä
        required_count = 50 # Määritä vaadittu määrä
        if not question_ids or len(question_ids) < required_count:
            app.logger.error(f"Ei voitu aloittaa simulaatiota käyttäjälle {current_user.id}: Liian vähän kysymyksiä ({len(question_ids)}/{required_count}).")
            return jsonify({'error': f'Simulaation luonti epäonnistui: Tietokannassa ei ole tarpeeksi kysymyksiä (vaaditaan {required_count}).'}), 503 # Service Unavailable

        # Alusta simulaation tiedot sessioon
        start_time = datetime.now(timezone.utc)
        time_limit_seconds = 3600 # 60 minuuttia

        session['simulation'] = {
            'user_id': current_user.id,
            'question_ids': question_ids,
            'answers': [None] * len(question_ids), # Lista käyttäjän vastauksille (indeksit)
            'current_index': 0,
            'start_time': start_time.isoformat(),
            'time_limit': time_limit_seconds,
            'time_remaining': time_limit_seconds # Alkuperäinen aika
        }
        session.modified = True # Merkitse sessio muokatuksi

        app.logger.info(f"Uusi simulaatio aloitettu käyttäjälle {current_user.id}: {len(question_ids)} kysymystä, {time_limit_seconds // 60} min.")

        # Palauta tieto onnistumisesta ja ehkä ensimmäinen kysymys?
        # Tai client hakee ensimmäisen kysymyksen erikseen /api/simulation/question/0 kautta
        return jsonify({
            'success': True,
            'message': 'Simulaatio aloitettu.',
            'total_questions': len(question_ids),
            'time_limit': time_limit_seconds
            # 'first_question_id': question_ids[0] # Voit palauttaa tämän optimointina
        })

    except Exception as e:
        app.logger.error(f"Virhe simulaation aloituksessa käyttäjälle {current_user.id}: {e}", exc_info=True)
        return jsonify({'error': 'Virhe simulaation aloituksessa.'}), 500


@app.route('/api/simulation/question/<int:index>')
@login_required
# @premium_required
def get_simulation_question_api(index):
    """Hakee yhden kysymyksen simulaatiota varten indeksin perusteella."""
    # Tarkista aktiivinen simulaatio sessiosta
    if 'simulation' not in session or session['simulation'].get('user_id') != current_user.id:
        return jsonify({'error': 'Ei aktiivista simulaatiota löytynyt.'}), 404

    sim_session = session['simulation']
    question_ids = sim_session.get('question_ids', [])

    # Tarkista indeksin kelvollisuus
    if not (0 <= index < len(question_ids)):
        app.logger.warning(f"Virheellinen kysymysindeksi {index} haettu simulaatiossa (Kysymyksiä: {len(question_ids)})")
        return jsonify({'error': 'Virheellinen kysymysindeksi.'}), 400

    question_id = question_ids[index]
    # Hae kysymys tietokannasta (ilman käyttäjän progressia tässä vaiheessa)
    question = db_manager.get_question_by_id(question_id, user_id=None)

    if question:
        # Päivitä sessioon nykyinen indeksi, jotta tiedetään missä mennään
        session['simulation']['current_index'] = index
        session.modified = True

        # Palauta kysymyksen tiedot (ilman vastausta/selitystä)
        q_dict = asdict(question)
        q_dict.pop('correct', None)
        q_dict.pop('explanation', None)
        q_dict.pop('question_normalized', None)
        # Palauta myös käyttäjän aiempi vastaus tähän kysymykseen, jos sellainen on
        q_dict['user_answer_index'] = sim_session.get('answers', [])[index]

        return jsonify(q_dict)
    else:
        app.logger.error(f"Simulaation kysymystä ID {question_id} (indeksi {index}) ei löytynyt tietokannasta!")
        # Poista virheellinen ID sessiosta? Tai palauta virhe.
        return jsonify({'error': f'Kysymystä ei löytynyt (ID: {question_id}).'}), 404


@app.route('/api/simulation/update', methods=['POST'])
@login_required
# @premium_required
def update_simulation_api():
    """
    Vastaanottaa päivityksen simulaation tilasta (vastaukset, aika) ja tallentaa sen sessioon.
    Tätä kutsutaan säännöllisesti clientilta (esim. 30s välein tai vastausta vaihdettaessa).
    """
    if 'simulation' not in session or session['simulation'].get('user_id') != current_user.id:
        return jsonify({'error': 'Ei aktiivista simulaatiota.'}), 404

    data = request.get_json()
    if not data:
        return jsonify({'error': 'Pyynnön body puuttuu.'}), 400

    sim_session = session['simulation']
    updated = False

    # Päivitä vastaukset, jos ne on lähetetty
    if 'answers' in data and isinstance(data['answers'], list):
        # Tässä voisi validoida, että vastausten määrä täsmää
        if len(data['answers']) == len(sim_session.get('question_ids',[])):
            sim_session['answers'] = data['answers']
            updated = True
        else:
             app.logger.warning(f"Simulaation päivitys: Vastausten määrä ({len(data['answers'])}) ei täsmää kysymysten määrään ({len(sim_session.get('question_ids',[]))}).")
             # Älä päivitä vastauksia, jos määrä ei täsmää

    # Päivitä jäljellä oleva aika
    if 'time_remaining' in data:
        try:
            time_remaining = int(data['time_remaining'])
            # Varmista, että aika ei ole negatiivinen
            sim_session['time_remaining'] = max(0, time_remaining)
            updated = True
            app.logger.debug(f"Simulaation aika päivitetty: {sim_session['time_remaining']}s jäljellä.")
        except (ValueError, TypeError):
            app.logger.warning(f"Virheellinen 'time_remaining' arvo päivityksessä: {data['time_remaining']}")

    # Päivitä nykyinen indeksi (missä käyttäjä on käyttöliittymässä)
    if 'current_index' in data:
         try:
             current_index = int(data['current_index'])
             # Varmista, että indeksi on sallituissa rajoissa
             if 0 <= current_index < len(sim_session.get('question_ids', [])):
                 sim_session['current_index'] = current_index
                 updated = True
             else:
                  app.logger.warning(f"Virheellinen 'current_index' päivityksessä: {current_index}")
         except (ValueError, TypeError):
              app.logger.warning(f"Virheellinen 'current_index' tyyppi päivityksessä: {data['current_index']}")


    # Tallenna muutokset sessioon vain jos jotain päivitettiin
    if updated:
        session.modified = True
        return jsonify({'success': True, 'message': 'Simulaation tila tallennettu.'})
    else:
        # Mitään ei päivitetty tai data oli virheellistä
        return jsonify({'success': False, 'message': 'Ei päivitettävää dataa tai data virheellistä.'}), 200 # Ei välttämättä virhe


@app.route('/api/simulation/submit', methods=['POST'])
@login_required
# @premium_required
def submit_simulation_api():
    """Vastaanottaa lopulliset vastaukset, laskee tuloksen, tallentaa sen ja tyhjentää session."""
    if 'simulation' not in session or session['simulation'].get('user_id') != current_user.id:
        return jsonify({'error': 'Ei aktiivista simulaatiota palautettavaksi.'}), 404

    sim_session = session['simulation']
    question_ids = sim_session.get('question_ids', [])
    user_answers_indices = sim_session.get('answers', []) # Haetaan sessioon tallennetut vastaukset

    # Laske kulunut aika
    try:
        start_time = datetime.fromisoformat(sim_session.get('start_time'))
        # Varmista aikavyöhyke (UTC)
        if start_time.tzinfo is None:
            start_time = start_time.replace(tzinfo=timezone.utc)
        end_time = datetime.now(timezone.utc)
        time_taken_seconds = int((end_time - start_time).total_seconds())
        # Voit myös käyttää clientin lähettämää aikaa tai sessioon tallennettua time_remaining
        time_limit = sim_session.get('time_limit', 3600)
        time_taken_seconds = min(time_taken_seconds, time_limit) # Aika ei voi ylittää rajaa
    except Exception as time_e:
        app.logger.error(f"Virhe simulaation ajan laskennassa: {time_e}")
        time_taken_seconds = sim_session.get('time_limit', 3600) # Oleta maksimiaika virheessä


    # Laske tulos ja kerää yksityiskohtaiset tulokset
    score = 0
    total_questions = len(question_ids)
    detailed_results = []
    answers_for_db = [] # JSON-tallennusta varten

    for i, question_id in enumerate(question_ids):
        question = db_manager.get_question_by_id(question_id, user_id=None) # Hae vain kysymysdata
        if not question:
            app.logger.error(f"Simulaation kysymystä ID {question_id} ei löytynyt tuloksia laskiessa!")
            # Voit päättää, lasketaanko tämä virheeksi vai ohitetaanko
            total_questions -= 1 # Tai älä vähennä ja laske virheeksi
            continue

        user_answer_index = user_answers_indices[i] if i < len(user_answers_indices) and user_answers_indices[i] is not None else None
        is_correct = (user_answer_index == question.correct)

        if is_correct:
            score += 1

        # Tallenna vastaus myös question_attempts tauluun (tilastointia varten)
        try:
             # Arvioi aika per kysymys (tarkempi aika vaatisi tallennusta clientilta)
             time_per_q = time_taken_seconds / len(question_ids) if len(question_ids) > 0 else 30
             db_manager.update_question_stats(question_id, is_correct, time_per_q, current_user.id)
        except Exception as e:
             app.logger.error(f"Virhe simulaatiovastauksen tallentamisessa attempts-tauluun (QID {question_id}): {e}")


        # Kerää tiedot tulosnäkymää varten
        user_answer_text = question.options[user_answer_index] if user_answer_index is not None and 0 <= user_answer_index < len(question.options) else "Ei vastausta"
        correct_answer_text = question.options[question.correct] if 0 <= question.correct < len(question.options) else "N/A"

        detailed_results.append({
            'question_id': question.id,
            'question_text': question.question,
            'options': question.options,
            'user_answer_index': user_answer_index,
            'user_answer_text': user_answer_text,
            'correct_answer_index': question.correct,
            'correct_answer_text': correct_answer_text,
            'is_correct': is_correct,
            'explanation': question.explanation
        })
        # Kerää tiedot JSON-tallennukseen (tiiviimpi muoto)
        answers_for_db.append({
            'q': question.id,
            's': user_answer_index, # selected
            'c': is_correct # correct
        })

    # Laske prosentti ja läpäisy
    percentage = round((score / total_questions * 100), 1) if total_questions > 0 else 0.0
    passed = percentage >= 80.0 # Määritä läpäisyraja

    # Tallenna simulaation tulos tietokantaan (tarvitaan uusi taulu/metodi)
    try:
        # Oletetaan, että on simulation_results-taulu ja metodi:
        # db_manager.save_simulation_result(
        #     user_id=current_user.id,
        #     score=score,
        #     total_questions=total_questions,
        #     time_taken=time_taken_seconds,
        #     passed=passed,
        #     answers_json=json.dumps(answers_for_db)
        # )
        app.logger.info(f"Simulaation tulos tallennettu käyttäjälle {current_user.id}: {score}/{total_questions} ({percentage}%), Läpäisty: {passed}")
        # Tarkista simulaatioon liittyvät saavutukset
        # achievement_manager.check_achievements(current_user.id, context={'simulation_score': percentage, 'simulation_passed': passed})

    except Exception as save_e:
        app.logger.error(f"Virhe simulaation tuloksen tallennuksessa: {save_e}", exc_info=True)
        # Älä estä tulosten näyttämistä clientille

    # Tyhjennä simulaatio sessiosta
    session.pop('simulation', None)
    session.modified = True

    # Palauta tulokset clientille
    return jsonify({
        'success': True,
        'score': score,
        'total_questions': total_questions,
        'percentage': percentage,
        'time_taken': time_taken_seconds,
        'passed': passed,
        'detailed_results': detailed_results
    })

# ==============================================================================
# --- SIVUJEN REITIT ---
# ==============================================================================

# --- Juuri ja Julkiset Sivut ---
@app.route("/")
def index_route():
    """Ohjaa oikealle aloitussivulle roolin mukaan tai kirjautumiseen."""
    if current_user.is_authenticated:
        # Tarkista rooli ja ohjaa
        if hasattr(current_user, 'role'):
            if current_user.role == 'superuser':
                 return redirect(url_for('admin_route')) # admin_route hoitaa superuser-ohjauksen
            elif current_user.role == 'admin':
                 return redirect(url_for('admin_route'))
            else: # Oletusrooli 'user'
                 return redirect(url_for('dashboard_route'))
        else:
            # Jos roolia ei jostain syystä löydy, kirjaa ulos ja ohjaa loginiin
            app.logger.error(f"Kirjautuneella käyttäjällä {getattr(current_user,'id','N/A')} ei ole roolia. Kirjataan ulos.")
            logout_user()
            flash("Istuntovirhe. Kirjaudu uudelleen.", "danger")
            return redirect(url_for('login_route'))
    # Jos ei kirjautunut, ohjaa kirjautumissivulle
    return redirect(url_for('login_route'))

@app.route("/privacy")
def privacy_route():
    """Näyttää tietosuojaselostesivun."""
    # Varmista, että templates/privacy.html on olemassa
    return render_template("privacy.html")

@app.route("/terms")
def terms_route():
    """Näyttää käyttöehdot-sivun."""
    # Varmista, että templates/terms.html on olemassa
    return render_template("terms.html")

# --- Kirjautuminen, Rekisteröinti, Uloskirjautuminen ---
# (Nämä ovat jo osiossa 2, Flask-Login ja Apufunktiot)
# Varmista, että ne ovat tässä tiedostossa vain kerran.

# --- Salasanan Palautus ---
# (Nämä ovat jo osiossa 2, Flask-Login ja Apufunktiot)
# Varmista, että ne ovat tässä tiedostossa vain kerran.


# --- Käyttäjän Perusreitit (vaativat kirjautumisen) ---

@app.route("/dashboard")
@login_required # Vaatii kirjautumisen (kaikki roolit pääsevät)
def dashboard_route():
    """Näyttää käyttäjän henkilökohtaisen dashboardin."""
    try:
        # Hae analytiikkatiedot
        analytics = stats_manager.get_learning_analytics(current_user.id)

        # Etsi suositukset (coach pick, strength pick)
        coach_pick = None
        strength_pick = None
        # Olettaen, että analytics['categories'] on lista dictejä, joissa 'success_rate' ja 'attempts'
        valid_categories = [cat for cat in analytics.get('categories', []) if cat.get('attempts', 0) >= 5 and cat.get('success_rate') is not None]
        if valid_categories:
            coach_pick = min(valid_categories, key=lambda x: x.get('success_rate', 101.0)) # 101, jotta None menee loppuun
            # Vahvuus vaatii esim. väh. 10 yritystä ja korkean onnistumisen
            strong_candidates = [cat for cat in valid_categories if cat.get('attempts', 0) >= 10]
            if strong_candidates:
                strength_pick = max(strong_candidates, key=lambda x: x.get('success_rate', -1.0)) # -1, jotta None menee alkuun

        # Hae kuitattavien kehityskohteiden määrä
        false_val = False if db_manager.is_postgres else 0
        mistake_count_res = db_manager._execute("""
            SELECT COUNT(*) as count FROM user_question_progress
            WHERE user_id = ? AND times_correct < times_shown
              AND (mistake_acknowledged IS NULL OR mistake_acknowledged = ?)
        """, (current_user.id, false_val), fetch='one')
        mistake_count = mistake_count_res['count'] if mistake_count_res else 0

        # Hae käyttäjän viimeisimmät harjoitteluasetukset (jos käytössä)
        # Nämä ovat nyt User-oliossa, jos load_user lataa ne
        # user_prefs = db_manager.get_user_by_id(current_user.id) # Tai käytä current_user suoraan
        last_categories_json = getattr(current_user, 'last_practice_categories', '[]')
        last_difficulties_json = getattr(current_user, 'last_practice_difficulties', '[]')
        try:
            last_categories = json.loads(last_categories_json or '[]')
            last_difficulties = json.loads(last_difficulties_json or '[]')
        except json.JSONDecodeError:
            last_categories = []
            last_difficulties = []

        # Hae kaikki kategoriat valintaa varten
        all_categories_from_db = db_manager.get_categories()

        # Tarkista, onko simulaatio kesken (Flaskin sessiosta)
        has_active_simulation = ('simulation' in session and session['simulation'].get('user_id') == current_user.id)

        return render_template(
            'dashboard.html',
            analytics=analytics, # Välitä koko analytics-dict templatelle
            coach_pick=coach_pick,
            strength_pick=strength_pick,
            mistake_count=mistake_count,
            categories=all_categories_from_db,
            last_categories=last_categories,
            last_difficulties=last_difficulties,
            has_active_simulation=has_active_simulation
        )
    except Exception as e:
        app.logger.error(f"Virhe dashboardin latauksessa käyttäjälle {current_user.id}: {e}", exc_info=True)
        flash("Virhe näkymän latauksessa. Yritä myöhemmin uudelleen.", "danger")
        # Yritä ohjata johonkin turvalliseen paikkaan, esim. profiiliin
        return redirect(url_for('profile_route'))


@app.route("/practice")
@login_required
def practice_route():
    """Näyttää yleisen harjoittelusivun."""
    # Välitä kategoriat valintaa varten
    categories = db_manager.get_categories()
    return render_template("practice.html", categories=categories)

@app.route("/practice/<category>")
@login_required
def practice_category_route(category):
    """Näyttää harjoittelusivun esivalitulla kategorialla."""
    # Varmista, että kategoria on validi (turvallisuussyistä)
    all_categories = db_manager.get_categories()
    if category not in all_categories and category != "Kaikki kategoriat":
        flash(f"Tuntematon kategoria: {category}", "warning")
        return redirect(url_for('practice_route'))

    return render_template("practice.html", selected_category=category, categories=all_categories)


@app.route("/review")
@login_required
def review_route():
    """Näyttää älykkään kertauksen sivun."""
    # Itse kysymykset haetaan API:n kautta (/api/review-questions)
    return render_template("review.html")

@app.route("/stats")
@login_required
def stats_route():
    """Näyttää käyttäjän tilastosivun."""
    # Tilastot haetaan API:n kautta (/api/stats)
    return render_template("stats.html")

@app.route("/achievements")
@login_required
def achievements_route():
    """Näyttää käyttäjän saavutussivun."""
    # Saavutukset haetaan API:n kautta (/api/achievements)
    return render_template("achievements.html")

@app.route("/mistakes")
@login_required
def mistakes_route():
    """Näyttää käyttäjän kehityskohteet (väärin vastatut kysymykset)."""
    # Kysymykset haetaan API:n kautta (/api/incorrect_questions)
    return render_template("mistakes.html")

@app.route("/calculator")
@login_required
def calculator_route():
    """Näyttää lääkelaskurin sivun."""
    # Varmista, että templates/calculator.html on olemassa
    return render_template("calculator.html")

# --- Simulaatiosivu ---
# Tämä käyttää nyt Flaskin sessiota tilan hallintaan
@app.route('/simulation')
@login_required
# @premium_required # Lisää tämä, jos simulaatio on maksullinen
def simulation_route():
    """Renderöi koesimulaatiosivun tai ohjaa aloittamaan uuden."""

    # Tarkista, onko pyyntö aloittaa uusi simulaatio
    if request.args.get('new') == 'true':
        # Poista vanha sessio, jos sellainen on
        if 'simulation' in session and session['simulation'].get('user_id') == current_user.id:
            session.pop('simulation', None)
            session.modified = True
            app.logger.info(f"Aloitetaan uusi simulaatio, vanha poistettu käyttäjälle {current_user.id}")

        # Ohjaa API-endpointtiin, joka HAKEE kysymykset ja LUO session
        # Tämä on parempi kuin kysymysten haku suoraan tässä reitissä
        # Client-puoli voi sitten kutsua /api/simulation/start ja ohjata tähän sivuun
        # TAI voit tehdä sen tässä:
        try:
             # Hae 50 satunnaista kysymys-ID:tä
             question_ids = db_manager.get_random_question_ids(limit=50)
             required_count = 50
             if not question_ids or len(question_ids) < required_count:
                  flash(f"Simulaation luonti epäonnistui: Tietokannassa ei ole tarpeeksi kysymyksiä (vaaditaan {required_count}).", "danger")
                  return redirect(url_for('dashboard_route'))

             # Alusta simulaation tiedot sessioon
             start_time = datetime.now(timezone.utc)
             time_limit_seconds = 3600 # 60 minuuttia
             session['simulation'] = {
                 'user_id': current_user.id, 'question_ids': question_ids,
                 'answers': [None] * len(question_ids), 'current_index': 0,
                 'start_time': start_time.isoformat(), 'time_limit': time_limit_seconds,
                 'time_remaining': time_limit_seconds
             }
             session.modified = True
             app.logger.info(f"Uusi simulaatio aloitettu (reitistä): Käyttäjä {current_user.id}")
             # Ohjaa sivuun, joka näyttää simulaation käyttöliittymän
             return redirect(url_for('simulation_route', resume='true')) # Lisätään resume-parametri

        except Exception as e:
             app.logger.error(f"Virhe uuden simulaation aloituksessa reitistä: {e}", exc_info=True)
             flash("Virhe simulaation aloituksessa.", "danger")
             return redirect(url_for('dashboard_route'))


    # Tarkista, onko olemassa oleva simulaatio sessiossa
    has_existing_session = ('simulation' in session and session['simulation'].get('user_id') == current_user.id)
    session_info = {} # Lähetetään templatelle tiedot näytettäväksi

    if has_existing_session:
        sim = session['simulation']
        # Laske jäljellä oleva aika (jos 'time_remaining' ei ole ajan tasalla)
        try:
            stored_time = sim.get('time_remaining')
            start_time_iso = sim.get('start_time')
            time_limit = sim.get('time_limit', 3600)

            if stored_time is None and start_time_iso: # Jos aikaa ei ole tallennettu, laske se alusta
                start_time = datetime.fromisoformat(start_time_iso)
                if start_time.tzinfo is None: start_time = start_time.replace(tzinfo=timezone.utc)
                elapsed = (datetime.now(timezone.utc) - start_time).total_seconds()
                time_remaining = max(0, int(time_limit - elapsed))
            elif stored_time is not None:
                time_remaining = max(0, int(stored_time))
            else:
                 time_remaining = time_limit # Oletus, jos kumpaakaan ei löydy

            # Päivitä sessioon varmuuden vuoksi
            sim['time_remaining'] = time_remaining
            session.modified = True

            # Kerää tiedot näkymää varten
            total_q = len(sim.get('question_ids', []))
            answered_q = len([a for a in sim.get('answers', []) if a is not None])
            session_info = {
                'current_index': sim.get('current_index', 0),
                'total_questions': total_q,
                'answered_questions': answered_q,
                'time_remaining_seconds': time_remaining
            }
            app.logger.info(f"Jatketaan simulaatiota: Käyttäjä {current_user.id}, Indeksi {session_info['current_index']}, Aikaa {time_remaining}s")

        except Exception as e:
            app.logger.error(f"Virhe simulaation tilan laskennassa: {e}", exc_info=True)
            # Nollaa sessio virhetilanteessa?
            session.pop('simulation', None)
            session.modified = True
            flash("Virhe simulaation latauksessa. Voit aloittaa uuden.", "warning")
            return redirect(url_for('dashboard_route'))

    # Renderöi simulaatiosivu
    # Varmista, että templates/simulation.html on olemassa
    # Tämä template hakee kysymykset API:n kautta dynaamisesti
    return render_template('simulation.html',
                          has_existing_session=has_existing_session,
                          session_info=session_info)


@app.route("/profile")
@login_required
def profile_route():
    """Näyttää käyttäjän profiilisivun."""
    try:
        # Hae perustilastot näytettäväksi
        analytics = stats_manager.get_learning_analytics(current_user.id)
        # Varmista, että templates/profile.html on olemassa
        return render_template("profile.html", stats=analytics.get('general', {}))
    except Exception as e:
        app.logger.error(f"Virhe profiilisivun latauksessa käyttäjälle {current_user.id}: {e}", exc_info=True)
        flash("Virhe profiilin latauksessa.", "danger")
        return redirect(url_for('dashboard_route'))


@app.route("/settings", methods=['GET', 'POST'])
@login_required
def settings_route():
    """Näyttää asetussivun ja käsittelee salasanan vaihdon."""
    if request.method == 'POST':
        # --- Salasanan vaihto ---
        current_password = request.form.get('current_password')
        new_password = request.form.get('new_password')
        confirm_password = request.form.get('confirm_password')

        # Tarkistukset
        error = False
        if not all([current_password, new_password, confirm_password]):
            flash('Täytä kaikki salasanakentät.', 'danger')
            error = True
        if new_password != confirm_password:
            flash('Uudet salasanat eivät täsmää.', 'danger')
            error = True
        # Lisää salasanan vahvuuden tarkistus uudelle salasanalle
        if len(new_password) < 8 or not re.search(r'[A-Z]', new_password) or not re.search(r'[a-z]', new_password) or not re.search(r'[0-9]', new_password):
             flash('Uuden salasanan tulee olla vähintään 8 merkkiä ja sisältää iso kirjain, pieni kirjain ja numero.', 'danger')
             error = True

        if not error:
            try:
                # Hae nykyinen hash tietokannasta
                user_data = db_manager.get_user_by_id(current_user.id)
                if user_data and bcrypt.check_password_hash(user_data['password'], current_password):
                    # Hashaa uusi salasana
                    new_hashed_password = bcrypt.generate_password_hash(new_password).decode('utf-8')
                    # Päivitä tietokantaan
                    success, db_error = db_manager.update_user_password(current_user.id, new_hashed_password)
                    if success:
                        flash('Salasana vaihdettu onnistuneesti!', 'success')
                        app.logger.info(f"Käyttäjä {current_user.username} vaihtoi salasanansa.")
                        # Kirjaa käyttäjä ulos turvallisuussyistä salasanan vaihdon jälkeen? (Valinnainen)
                        # logout_user()
                        # return redirect(url_for('login_route'))
                    else:
                        flash(f'Salasanan päivitys epäonnistui: {db_error}', 'danger')
                        app.logger.error(f"Salasanan päivitys epäonnistui DB-tasolla käyttäjälle {current_user.id}: {db_error}")
                else:
                    flash('Nykyinen salasana on väärä.', 'danger')
                    app.logger.warning(f"Väärä nykyinen salasana syötetty käyttäjälle {current_user.username}")
            except Exception as e:
                flash('Salasanan vaihdossa tapahtui odottamaton virhe.', 'danger')
                app.logger.error(f"Virhe salasanan vaihdossa käyttäjälle {current_user.id}: {e}", exc_info=True)

        # Ohjaa takaisin asetussivulle näyttämään flash-viestit
        return redirect(url_for('settings_route'))

    # GET-pyyntö: Näytä asetussivu
    # Hae käyttäjän nykyiset asetukset tuoreeltaan tietokannasta
    user_settings = db_manager.get_user_by_id(current_user.id)
    # Varmista, että templates/settings.html on olemassa
    return render_template("settings.html", user_settings=user_settings)

# ==============================================================================
# --- ADMIN / SUPERUSER YLEINEN SISÄÄNKÄYNTI ---
# ==============================================================================

@app.route("/admin") # Tämä toimii nyt yleisenä admin-sisäänkäyntinä
@admin_required # Sallii 'admin' ja 'superuser' roolit
def admin_route():
    """
    Admin/Superuser-pääsivu. Ohjaa tai renderöi oikean dashboardin roolin mukaan.
    """
    try:
        # --- Superuser Näkymä ---
        if current_user.role == 'superuser':
            app.logger.debug(f"Ladataan Superuser Dashboard käyttäjälle {current_user.username}")
            organizations = db_manager.get_all_organizations() # Hakee aktiiviset orgit + käyttäjämäärät

            # Laske järjestelmän kokonaistilastot (vain aktiiviset käyttäjät)
            total_users_res = db_manager._execute(
                "SELECT COUNT(*) as count FROM users WHERE status='active'", fetch='one'
            )
            total_questions_res = db_manager._execute(
                "SELECT COUNT(*) as count FROM questions", fetch='one' # Kaikki kysymykset
            )
            total_validated_questions_res = db_manager._execute(
                "SELECT COUNT(*) as count FROM questions WHERE status='validated'", fetch='one' # Validoidut
            )

            stats = {
                'total_users': total_users_res['count'] if total_users_res else 0,
                'total_organizations': len(organizations) if organizations else 0,
                'total_questions': total_questions_res['count'] if total_questions_res else 0,
                'validated_questions': total_validated_questions_res['count'] if total_validated_questions_res else 0,
            }
            # Luo superuser_dashboard.html template tätä varten
            return render_template("superuser_dashboard.html", organizations=organizations, stats=stats)

        # --- Organisaation Admin Näkymä ('admin' rooli) ---
        elif current_user.role == 'admin':
            app.logger.debug(f"Ladataan Organisaatio Admin Dashboard käyttäjälle {current_user.username} (Org ID: {current_user.organization_id})")
            org_id = current_user.organization_id
            # organization_id on varmistettu @admin_required dekoraattorissa

            # Hae organisaation nimi
            org_data = db_manager._execute("SELECT name FROM organizations WHERE id = ?", (org_id,), fetch='one')
            org_name = org_data['name'] if org_data else "Tuntematon Organisaatio"

            # Hae organisaatiokohtaiset tilastot (vain aktiiviset käyttäjät)
            # Tarvitaan mahdollisesti uusia funktioita db_manageriin/stats_manageriin
            active_students_count = db_manager._execute(
                "SELECT COUNT(*) as count FROM users WHERE organization_id = ? AND role='user' AND status='active'",
                (org_id,), fetch='one'
            )['count'] or 0

            # Esim. keskimääräinen onnistumisprosentti TÄSSÄ organisaatiossa
            avg_success_rate_org = stats_manager.get_average_success_rate_for_organization(org_id) # Olettaa tämän metodin olemassaolon

            org_stats = {
                 'organization_name': org_name,
                 'active_students': active_students_count,
                 'avg_success_rate_org': f"{avg_success_rate_org:.1f}%" if avg_success_rate_org is not None else "N/A"
                 # Lisää muita relevantteja tilastoja tähän
            }

            # Hae myös yleisiä tietoja, kuten kysymysten kokonaismäärä
            total_questions_res = db_manager._execute("SELECT COUNT(*) as count FROM questions", fetch='one')
            general_stats = {
                'total_questions': total_questions_res['count'] if total_questions_res else 0
            }

            # Käytä vanhaa admin.html-templatea tai luo uusi org_admin_dashboard.html
            # Välitä sekä organisaatiokohtaiset että yleiset tiedot
            return render_template("admin.html", # Tai org_admin_dashboard.html
                                 org_stats=org_stats,
                                 general_stats=general_stats
                                 )
        else:
            # Tähän ei pitäisi päätyä @admin_required vuoksi
            app.logger.error(f"Pääsyvirhe /admin reitissä käyttäjälle {current_user.id} tuntemattomalla roolilla: {getattr(current_user, 'role', 'Ei roolia')}")
            flash("Luvaton pääsy.", "danger")
            return redirect(url_for('dashboard_route'))

    except Exception as e:
        flash(f'Virhe hallintanäkymän lataamisessa: {e}', 'danger')
        app.logger.error(f"Admin/Superuser page load error: {e}", exc_info=True)
        # Ohjaa turvallisesti käyttäjän omaan dashboardiin
        return redirect(url_for('dashboard_route'))

# ==============================================================================
# --- ORGANISAATION ADMIN-REITIT (Opettaja) ---
# ==============================================================================

@app.route("/admin/users")
@admin_required # Sallii admin ja superuser, mutta logiikka erottelee
def admin_users_route():
    """
    Näyttää organisaation adminille (opettajalle) oman organisaationsa käyttäjät.
    Superuser ohjataan omaan näkymäänsä.
    """
    # Ohjaa superuser pois tästä näkymästä
    if current_user.role == 'superuser':
        flash("Hallitse organisaatioita ja niiden käyttäjiä Superuser-näkymästä.", "info")
        return redirect(url_for('superuser_orgs_route'))

    # Tästä eteenpäin koodi koskee vain roolia 'admin' (opettaja)
    try:
        org_id = current_user.organization_id # Varmistettu jo dekoraattorissa
        app.logger.debug(f"Haetaan käyttäjiä organisaatiolle {org_id} adminille {current_user.username}")

        # Hae VAIN oman organisaation AKTIIVISET käyttäjät
        users = db_manager.get_users_for_organization(org_id)

        # Hae organisaation nimi näkymää varten
        org_data = db_manager._execute("SELECT name FROM organizations WHERE id = ?", (org_id,), fetch='one')
        org_name = org_data['name'] if org_data else "Oma Organisaatio"

        # Luo admin_users.html template tätä varten
        # Välitä myös organization_id, jotta lomakkeet toimivat oikein
        return render_template("admin_users.html",
                               users=users,
                               organization_name=org_name,
                               organization_id=org_id)

    except Exception as e:
        flash(f'Virhe käyttäjien haussa: {e}', 'danger')
        app.logger.error(f"Admin users fetch error for org {current_user.organization_id}: {e}", exc_info=True)
        return redirect(url_for('admin_route')) # Ohjaa admin-dashboardiin


@app.route("/admin/user/create", methods=['POST'])
@admin_required # Vain admin (opettaja) voi luoda käyttäjiä OMAAN organisaatioonsa
def admin_create_user_route():
    """Luo uuden käyttäjän (opiskelijan) adminin omaan organisaatioon."""
    # Varmista rooli uudelleen
    if current_user.role != 'admin':
         flash("Vain organisaation ylläpitäjä voi luoda käyttäjiä.", "danger")
         # Superuser ohjataan omaan reittiinsä, muut ulos
         if current_user.role == 'superuser': return redirect(url_for('admin_route'))
         return redirect(url_for('dashboard_route'))

    org_id = current_user.organization_id # Haetaan kirjautuneen adminin org ID

    # Haetaan tiedot lomakkeelta
    username = request.form.get('username','').strip()
    email = request.form.get('email','').strip().lower()
    # Rooli on aina 'user' kun admin luo käyttäjän
    role = 'user'
    expires_at_str = request.form.get('expires_at') # Oletetaan VVVV-KK-PP

    # Validoinnit
    error_redirect_url = url_for('admin_users_route') # Minne ohjataan virheessä
    if not username or not email:
        flash('Käyttäjänimi ja sähköposti ovat pakollisia.', 'danger')
        return redirect(error_redirect_url)
    if len(username) < 3:
         flash('Käyttäjänimen tulee olla vähintään 3 merkkiä.', 'danger')
         return redirect(error_redirect_url)
    # Lisää email regex-validointi
    email_regex = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    if not re.match(email_regex, email):
        flash('Virheellinen sähköpostiosoite.', 'danger')
        return redirect(error_redirect_url)


    # Käsittele vanhenemispäivä
    expires_at = None
    if expires_at_str:
        try:
            expires_at = datetime.strptime(expires_at_str, '%Y-%m-%d').replace(hour=23, minute=59, second=59)
            # Varmista, että vanhenemispäivä on tulevaisuudessa
            if expires_at < datetime.now():
                 flash("Vanhenemispäivämäärä ei voi olla menneisyydessä.", "danger")
                 return redirect(error_redirect_url)
        except ValueError:
            flash("Virheellinen vanhenemispäivämäärä. Käytä muotoa VVVV-KK-PP.", "danger")
            return redirect(error_redirect_url)

    # Luo salasana ja hashaa se
    password = generate_secure_password(10)
    try:
        hashed_password = bcrypt.generate_password_hash(password).decode('utf-8')
    except Exception as hash_e:
        flash(f"Salasanan luonti epäonnistui: {hash_e}", "danger")
        app.logger.error(f"Password hashing failed for admin create: {hash_e}")
        return redirect(error_redirect_url)

    # Yritä luoda käyttäjä tietokantaan
    success, error_msg = db_manager.create_user(
        username, email, hashed_password,
        role=role, # Aina 'user' tässä
        organization_id=org_id,
        expires_at=expires_at
    )

    if success:
        # ÄLÄ näytä salasanaa flash-viestissä tuotannossa!
        flash_message = f"Opiskelijatili '{username}' luotu onnistuneesti."
        app.logger.info(f"Admin {current_user.username} (Org ID: {org_id}) created user '{username}'")
        # TÄRKEÄÄ: Lähetä salasana turvallisesti käyttäjälle (esim. sähköpostilla tai näytä kerran)
        # Voit lisätä salasanan sessioon ja näyttää sen vain kerran seuraavalla sivulla:
        session['new_user_password'] = password
        session['new_user_username'] = username
        flash(flash_message, "success")
        # Ohjaa takaisin käyttäjälistaan, jossa salasana voidaan näyttää
        return redirect(url_for('admin_users_route', new_user=username))
    else:
        flash(f"Virhe käyttäjän luonnissa: {error_msg}", "danger")
        app.logger.error(f"Admin {current_user.username} failed to create user '{username}': {error_msg}")
        return redirect(error_redirect_url)


@app.route("/admin/deactivate_user/<int:user_id>", methods=['POST'])
@admin_required # Vain oman organisaation admin voi deaktivoida
def admin_deactivate_user_route(user_id):
    """Deaktivoi käyttäjän (Soft Delete)."""
    # Varmista rooli
    if current_user.role != 'admin':
         flash("Toiminto vaatii organisaation ylläpitäjän oikeudet.", "danger")
         return redirect(url_for('dashboard_route'))

    org_id = current_user.organization_id
    app.logger.info(f"Admin {current_user.username} (Org ID: {org_id}) yrittää deaktivoida käyttäjää {user_id}")

    try:
        # Hae poistettava käyttäjä ja varmista omistajuus
        user_to_deactivate = db_manager.get_user_by_id(user_id)

        if not user_to_deactivate:
            flash("Käyttäjää ei löytynyt.", "danger")
            return redirect(url_for('admin_users_route'))

        # Tarkista organisaatio
        if user_to_deactivate.get('organization_id') != org_id:
            app.logger.warning(f"Luvaton deaktivointiyritys: Admin {current_user.username} (Org ID: {org_id}) -> User {user_id} (Org ID: {user_to_deactivate.get('organization_id')})")
            flash("Toiminto kielletty. Käyttäjä ei kuulu organisaatioosi.", "danger")
            return redirect(url_for('admin_users_route'))

        # Estä itsensä deaktivointi
        if user_id == current_user.id:
            flash("Et voi deaktivoida itseäsi.", "warning")
            return redirect(url_for('admin_users_route'))

        # Suorita deaktivointi (soft delete)
        success, error = db_manager.deactivate_user(user_id)

        if success:
            flash(f'Käyttäjä #{user_id} ({user_to_deactivate.get("username", "N/A")}) deaktivoitu onnistuneesti.', 'success')
            app.logger.info(f"Admin {current_user.username} deactivated user {user_id} ('{user_to_deactivate.get('username', 'N/A')}')")
        else:
            flash(f'Virhe käyttäjän deaktivoinnissa: {error}', 'danger')
            app.logger.error(f"User deactivate error for ID {user_id} by admin {current_user.username}: {error}")

    except Exception as e:
        flash(f"Odottamaton virhe käyttäjän deaktivoinnissa: {e}", "danger")
        app.logger.error(f"Unexpected error during user deactivation {user_id}: {e}", exc_info=True)

    # Ohjaa aina takaisin käyttäjälistaan
    return redirect(url_for('admin_users_route'))

# app.py - LISÄÄ NÄMÄ REITIT

# ... (muiden admin-reittien, kuten /admin/deactivate_user, jälkeen) ...

@app.route("/admin/toggle_user_status/<int:user_id>", methods=['POST'])
@admin_required # Vain organisaation admin
def admin_toggle_user_status_route(user_id):
    """Vaihtaa käyttäjän statuksen (active/inactive) organisaation sisällä."""
    if current_user.role == 'superuser':
        flash("Superuser käyttää omaa näkymäänsä.", "warning")
        return redirect(url_for('admin_route')) # Ohjaa SU dashboardiin
    org_id = current_user.organization_id

    user_to_toggle = db_manager.get_user_by_id(user_id)

    # Turvatarkistukset
    if not user_to_toggle or user_to_toggle.get('organization_id') != org_id:
        flash("Käyttäjää ei löytynyt tai hän ei kuulu organisaatioosi.", "danger")
        return redirect(url_for('admin_users_route'))
    if user_id == current_user.id:
        flash("Et voi muuttaa omaa statustasi.", "warning")
        return redirect(url_for('admin_users_route'))
    # Voit estää muiden adminien statuksen muuttamisen, jos haluat:
    # if user_to_toggle.get('role') == 'admin':
    #     flash("Et voi muuttaa toisen adminin statusta.", "danger")
    #     return redirect(url_for('admin_users_route'))

    # Määritä uusi status POST-datasta tai laskemalla
    new_status_from_post = request.form.get('new_status') # Voi tulla JS:n kautta
    if new_status_from_post in ['active', 'inactive']:
         new_status = new_status_from_post
    else: # Laske oletuksena
        new_status = 'active' if user_to_toggle.get('status') == 'inactive' else 'inactive'

    # Päivitä tietokanta
    success, error = db_manager.update_user(user_id, {'status': new_status})

    if success:
        action_text = "aktivoitu" if new_status == 'active' else "deaktivoitu"
        flash(f'Käyttäjä {user_to_toggle.get("username","")} {action_text}.', 'success')
        app.logger.info(f"Admin {current_user.username} toggled status for user {user_id} to {new_status}")
    else:
        flash(f'Statuksen vaihto epäonnistui: {error}', 'danger')
        app.logger.error(f"Admin status toggle failed for user {user_id}: {error}")

    return redirect(url_for('admin_users_route'))


@app.route("/admin/toggle_user_role/<int:user_id>", methods=['POST'])
@admin_required # Vain organisaation admin
def admin_toggle_user_role_route(user_id):
    """Vaihtaa käyttäjän roolin (user <-> admin) organisaation sisällä."""
    if current_user.role == 'superuser':
        flash("Superuser käyttää omaa näkymäänsä.", "warning")
        return redirect(url_for('admin_route'))
    org_id = current_user.organization_id

    user_to_toggle = db_manager.get_user_by_id(user_id)

    # Turvatarkistukset
    if not user_to_toggle or user_to_toggle.get('organization_id') != org_id:
        flash("Käyttäjää ei löytynyt tai hän ei kuulu organisaatioosi.", "danger")
        return redirect(url_for('admin_users_route'))
    if user_id == current_user.id:
        flash("Et voi muuttaa omaa rooliasi.", "warning")
        return redirect(url_for('admin_users_route'))
    # Superuserin roolia ei voi muuttaa
    if user_to_toggle.get('role') == 'superuser':
         flash("Pääkäyttäjän roolia ei voi muuttaa.", "warning")
         return redirect(url_for('admin_users_route'))

    # Määritä uusi rooli POST-datasta tai laskemalla
    new_role_from_post = request.form.get('new_role')
    if new_role_from_post in ['user', 'admin']:
         new_role = new_role_from_post
    else: # Laske oletuksena
        current_role = user_to_toggle.get('role', 'user')
        new_role = 'admin' if current_role == 'user' else 'user'

    # Lisätarkistus: Varmista, että organisaatioon jää vähintään yksi admin?
    if user_to_toggle.get('role') == 'admin' and new_role == 'user':
        admin_count_res = db_manager._execute(
             "SELECT COUNT(*) as count FROM users WHERE organization_id = ? AND role='admin' AND status='active'",
             (org_id,), fetch='one'
        )
        admin_count = admin_count_res['count'] if admin_count_res else 0
        if admin_count <= 1:
            flash("Organisaatiossa on oltava vähintään yksi admin. Roolia ei voi poistaa.", "danger")
            return redirect(url_for('admin_users_route'))

    # Päivitä rooli tietokantaan
    success, error = db_manager.update_user_role(user_id, new_role) # Käytä tätä funktiota

    if success:
        action_text = "organisaation adminiksi" if new_role == 'admin' else "opiskelijaksi"
        flash(f'Käyttäjän {user_to_toggle.get("username","")} rooli muutettu: {action_text}.', 'success')
        app.logger.info(f"Admin {current_user.username} toggled role for user {user_id} to {new_role}")
    else:
        flash(f'Roolin vaihto epäonnistui: {error}', 'danger')
        app.logger.error(f"Admin role toggle failed for user {user_id}: {error}")

    return redirect(url_for('admin_users_route'))

# Lisää myös reitti Superuserin käyttäjän muokkaukselle (GET ja POST)
@app.route("/superuser/user/edit/<int:user_id>", methods=['GET', 'POST'])
@superuser_required
def superuser_edit_user_route(user_id):
    """Näyttää lomakkeen käyttäjän tietojen muokkaamiseen ja käsittelee sen."""
    user = db_manager.get_user_by_id(user_id)
    if not user:
        flash("Käyttäjää ei löytynyt.", "danger")
        return redirect(url_for('superuser_orgs_route'))

    # Estä itsensä tai ID 1 muokkaus joissain tapauksissa?
    # if user_id == current_user.id or user_id == 1: ...

    if request.method == 'POST':
        # Kerää tiedot lomakkeelta
        data_to_update = {}
        data_to_update['username'] = request.form.get('username','').strip()
        data_to_update['email'] = request.form.get('email','').strip().lower()
        new_role = request.form.get('role')
        new_org_id_str = request.form.get('organization_id')
        expires_at_str = request.form.get('expires_at')

        # Validoi ja päivitä tiedot
        error = False
        if not data_to_update['username'] or not data_to_update['email']:
            flash("Käyttäjänimi ja sähköposti ovat pakollisia.", "danger")
            error = True
        # Lisää validointeja...

        # Käsittele vanhenemispäivä
        expires_at = None
        if expires_at_str:
            try:
                expires_at = datetime.strptime(expires_at_str, '%Y-%m-%d').replace(hour=23, minute=59, second=59)
                if expires_at < datetime.now().replace(hour=0, minute=0, second=0, microsecond=0) and not error: # Salli vanhentuneen päivän asetus, jos ei muita virheitä
                    # Voit näyttää varoituksen, jos asetetaan menneisyyteen
                     pass
            except ValueError:
                flash("Virheellinen vanhenemispäivämäärä.", "danger")
                error = True
        data_to_update['expires_at'] = expires_at

        # Käsittele organisaatio
        try:
            # Salli organisaation poisto asettamalla se tyhjäksi merkkijonoksi tai 0:ksi lomakkeella
            if new_org_id_str and new_org_id_str != 'None' and new_org_id_str != '0':
                data_to_update['organization_id'] = int(new_org_id_str)
                # Tarkista, että organisaatio on olemassa
                org_check = db_manager._execute("SELECT id FROM organizations WHERE id = ?", (data_to_update['organization_id'],), fetch='one')
                if not org_check:
                     flash("Valittua organisaatiota ei löydy.", "danger")
                     error = True
            else:
                 data_to_update['organization_id'] = None # Aseta NULL, jos tyhjä tai 'None'
        except ValueError:
            flash("Virheellinen organisaatio ID.", "danger")
            error = True

        if error:
             # Palauta lomake täytetyillä (mutta virheellisillä) tiedoilla
             organizations = db_manager.get_all_organizations() # Hae organisaatiot uudelleen
             return render_template('superuser_edit_user.html', user=request.form, organizations=organizations, current_org_id=request.form.get('organization_id')), 400


        # Päivitä perustiedot (username, email, expires, org)
        success_basic, error_basic = db_manager.update_user(user_id, data_to_update)

        # Päivitä rooli erikseen (jos se muuttui ja on sallittu)
        success_role, error_role = True, None
        original_role = user.get('role')
        if new_role and new_role != original_role and new_role in ['user', 'admin']:
             # Lisää tarkistus: Älä anna muuttaa viimeistä adminia useriksi?
             # Lisää tarkistus: Älä anna muuttaa superuserin roolia?
             if user.get('role') == 'superuser':
                  flash("Superuserin roolia ei voi muuttaa.", "danger")
                  success_role = False
                  error_role = "Superuserin roolia ei voi muuttaa."
             else:
                  # Varmista, että adminilla on organisaatio
                  if new_role == 'admin' and data_to_update.get('organization_id') is None:
                       flash("Admin-rooli vaatii organisaation.", "danger")
                       success_role = False
                       error_role = "Admin-rooli vaatii organisaation."
                  else:
                       success_role, error_role = db_manager.update_user_role(user_id, new_role)


        if success_basic and success_role:
            flash(f'Käyttäjän {data_to_update["username"]} tiedot päivitetty.', 'success')
            app.logger.info(f"Superuser {current_user.username} updated user {user_id}")
            # Ohjaa takaisin organisaationäkymään
            org_id = data_to_update.get('organization_id')
            return redirect(url_for('superuser_view_organization_route', org_id=org_id) if org_id else url_for('superuser_orgs_route'))
        else:
            # Kerää virheet yhteen
            errors = []
            if error_basic: errors.append(f"Perustietojen päivitys: {error_basic}")
            if error_role: errors.append(f"Roolin päivitys: {error_role}")
            flash(f'Päivitys epäonnistui: {"; ".join(errors)}', 'danger')
            app.logger.error(f"Superuser update failed for user {user_id}: Basic: {error_basic}, Role: {error_role}")
            # Palauta lomake täytetyillä tiedoilla
            organizations = db_manager.get_all_organizations()
            # Yhdistä vanhat ja uudet tiedot näyttämistä varten
            user_form_data = {**user, **request.form}
            return render_template('superuser_edit_user.html', user=user_form_data, organizations=organizations, current_org_id=user_form_data.get('organization_id')), 500


    # GET-pyyntö: Näytä lomake
    organizations = db_manager.get_all_organizations()
    # Luo superuser_edit_user.html template
    return render_template('superuser_edit_user.html', user=user, organizations=organizations, current_org_id=user.get('organization_id'))

@app.route('/admin/edit_user_settings/<int:user_id>', methods=['POST'])
@admin_required # Vain admin voi muokata oman organisaationsa käyttäjien asetuksia
def edit_user_settings(user_id):
    """Päivittää käyttäjän asetuksia (esim. häiriötekijät)."""
     # Varmista rooli
    if current_user.role != 'admin':
         # Superuser voi käyttää omaa mekanismiaan tai tätä ei sallita heille
         flash("Toiminto ei sallittu roolillesi.", "danger")
         return redirect(url_for('dashboard_route'))

    org_id = current_user.organization_id

    # Hae muokattava käyttäjä ja varmista omistajuus
    user_to_edit = db_manager.get_user_by_id(user_id)
    if not user_to_edit or user_to_edit.get('organization_id') != org_id:
        flash("Käyttäjää ei löytynyt tai hän ei kuulu organisaatioosi.", "danger")
        return redirect(url_for('admin_users_route'))

    # Kerää päivitettävät tiedot lomakkeelta
    try:
        data_to_update = {
            # Muunna checkboxin arvo ('on'/'None' tai '1'/'0') booleaniksi
            'distractors_enabled': request.form.get('distractors_enabled') == 'on' or request.form.get('distractors_enabled') == '1',
            # Varmista integer ja rajoita arvo
            'distractor_probability': max(0, min(100, int(request.form.get('distractor_probability', 25))))
        }
    except ValueError:
         flash("Virheellinen syöte todennäköisyydelle.", "danger")
         return redirect(url_for('admin_users_route'))


    # Suorita päivitys
    success, error = db_manager.update_user(user_id, data_to_update)

    if success:
        flash(f'Käyttäjän {user_to_edit["username"]} asetukset päivitetty onnistuneesti!', 'success')
        app.logger.info(f"Admin {current_user.username} updated settings for user {user_id}: {data_to_update}")
    else:
        flash(f'Virhe päivitettäessä asetuksia: {error}', 'danger')
        app.logger.error(f"Admin {current_user.username} failed to update settings for user {user_id}: {error}")

    return redirect(url_for('admin_users_route'))


# --- Globaalien kysymysten hallinta (VAIN Superuser) ---
# Siirretään nämä Superuser-osioon ja varmistetaan @superuser_required

# (Siirrä /admin/questions, /admin/add_question, /admin/edit_question,
#  /admin/delete_question, /admin/bulk_upload, /admin/find_duplicates,
#  /admin/clear_database reitit tänne Superuser-osioon ja vaihda dekoraattori)

# ==============================================================================
# --- SUPERUSER-KOHTAISET REITIT ---
# ==============================================================================

@app.route("/superuser/organizations")
@superuser_required
def superuser_orgs_route():
    """Näyttää kaikki organisaatiot ja lomakkeen uuden luomiseen."""
    try:
        # Hakee aktiiviset organisaatiot ja niiden aktiivisten käyttäjien määrän
        organizations = db_manager.get_all_organizations()
        # Luo superuser_organizations.html template tätä varten
        return render_template("superuser_organizations.html", organizations=organizations)
    except Exception as e:
        flash(f"Virhe organisaatioiden haussa: {e}", "danger")
        app.logger.error(f"Superuser org fetch error: {e}", exc_info=True)
        # Ohjaa superuser dashboardiin (admin_route hoitaa ohjauksen)
        return redirect(url_for('admin_route'))

@app.route("/superuser/organization/create", methods=['POST'])
@superuser_required
def superuser_create_org_route():
    """Luo uuden organisaation."""
    name = request.form.get('name','').strip()
    contact_person = request.form.get('contact_person','').strip() or None
    contact_email = request.form.get('contact_email','').strip().lower() or None

    if not name:
        flash("Organisaation nimi on pakollinen.", "danger")
    else:
        success, error_msg = db_manager.create_organization(name, contact_person, contact_email)
        if success:
            flash(f"Organisaatio '{name}' luotu onnistuneesti.", "success")
            app.logger.info(f"Superuser {current_user.username} created organization '{name}'")
        else:
            flash(f"Virhe organisaation luonnissa: {error_msg}", "danger")
            app.logger.error(f"Superuser org creation failed for '{name}': {error_msg}")

    return redirect(url_for('superuser_orgs_route'))

@app.route("/superuser/organization/<int:org_id>")
@superuser_required
def superuser_view_organization_route(org_id):
     """Näyttää yhden organisaation tiedot ja käyttäjälistan (korvaa vanhan /users reitin)."""
     try:
        organization = db_manager._execute("SELECT * FROM organizations WHERE id = ?", (org_id,), fetch='one')
        if not organization:
            flash("Organisaatiota ei löytynyt.", "danger")
            return redirect(url_for('superuser_orgs_route'))

        # Hae KAIKKI käyttäjät (aktiiviset ja inaktiiviset) tästä organisaatiosta
        all_users_in_org = db_manager._execute(
             """SELECT id, username, email, role, status, created_at, expires_at
                FROM users WHERE organization_id = ? ORDER BY status DESC, username ASC""",
             (org_id,), fetch='all'
        )

        # Luo superuser_org_users.html tai vastaava template
        return render_template("superuser_view_organization.html",
                               users=all_users_in_org,
                               organization=organization)
     except Exception as e:
        flash(f"Virhe organisaation tietojen haussa: {e}", "danger")
        app.logger.error(f"Superuser view org error for org {org_id}: {e}", exc_info=True)
        return redirect(url_for('superuser_orgs_route'))


@app.route("/superuser/user/create", methods=['POST'])
@superuser_required
def superuser_create_user_route():
    """Luo uuden käyttäjän (user tai admin) määritettyyn organisaatioon."""
    # Haetaan tiedot lomakkeelta
    username = request.form.get('username','').strip()
    email = request.form.get('email','').strip().lower()
    role = request.form.get('role', 'user')
    organization_id_str = request.form.get('organization_id')
    expires_at_str = request.form.get('expires_at')

    # Validoi organisaatio ID
    try:
        organization_id = int(organization_id_str)
        # Varmista, että organisaatio on olemassa
        org_check = db_manager._execute("SELECT id FROM organizations WHERE id = ?", (organization_id,), fetch='one')
        if not org_check:
             raise ValueError("Organisaatiota ei löydy")
    except (ValueError, TypeError):
        flash("Virheellinen tai puuttuva organisaatio ID.", "danger")
        # Yritä ohjata takaisin, jos mahdollista, muuten yleisnäkymään
        redirect_url = request.referrer or url_for('superuser_orgs_route')
        return redirect(redirect_url)

    # Muut validoinnit
    if not username or not email:
        flash("Käyttäjänimi ja sähköposti ovat pakollisia.", "danger")
        return redirect(url_for('superuser_view_organization_route', org_id=organization_id))
    if role not in ['user', 'admin']: # Superuser ei voi luoda toista superuseria tätä kautta
        flash("Virheellinen rooli valittu (sallittu: user, admin).", "danger")
        return redirect(url_for('superuser_view_organization_route', org_id=organization_id))
    # Lisää email ja username validoinnit (regex, pituus)

    # Käsittele vanhenemispäivä
    expires_at = None
    if expires_at_str:
        try:
            expires_at = datetime.strptime(expires_at_str, '%Y-%m-%d').replace(hour=23, minute=59, second=59)
            if expires_at < datetime.now():
                 flash("Vanhenemispäivämäärä ei voi olla menneisyydessä.", "danger")
                 return redirect(url_for('superuser_view_organization_route', org_id=organization_id))
        except ValueError:
            flash("Virheellinen vanhenemispäivämäärä. Käytä muotoa VVVV-KK-PP.", "danger")
            return redirect(url_for('superuser_view_organization_route', org_id=organization_id))

    # Luo salasana ja hashaa se
    password = generate_secure_password(10)
    try:
        hashed_password = bcrypt.generate_password_hash(password).decode('utf-8')
    except Exception as hash_e:
        flash(f"Salasanan luonti epäonnistui: {hash_e}", "danger")
        app.logger.error(f"Password hashing failed for superuser create: {hash_e}")
        return redirect(url_for('superuser_view_organization_route', org_id=organization_id))

    # Yritä luoda käyttäjä
    success, error_msg = db_manager.create_user(
        username, email, hashed_password,
        role=role,
        organization_id=organization_id,
        expires_at=expires_at
    )

    if success:
        flash_message = f"Käyttäjä '{username}' ({role}) luotu organisaatioon ID {organization_id}."
        if DEBUG_MODE: flash_message += f" Salasana (DEV): {password}"
        flash(flash_message, "success")
        app.logger.info(f"Superuser {current_user.username} created user '{username}' ({role}) for org {organization_id}")
        # Tässä voisi lähettää tervetuloviestin uudelle käyttäjälle
    else:
        flash(f"Virhe käyttäjän luonnissa: {error_msg}", "danger")
        app.logger.error(f"Superuser user creation failed for '{username}': {error_msg}")

    return redirect(url_for('superuser_view_organization_route', org_id=organization_id))


@app.route("/superuser/reactivate_user/<int:user_id>", methods=['POST'])
@superuser_required
def superuser_reactivate_user_route(user_id):
    """Aktivoi 'inactive'-tilassa olevan käyttäjän uudelleen."""
    # Hae käyttäjä ensin
    user_to_reactivate = db_manager.get_user_by_id(user_id)
    if not user_to_reactivate:
        flash("Käyttäjää ei löytynyt.", "danger")
        return redirect(request.referrer or url_for('superuser_orgs_route'))

    org_id = user_to_reactivate.get('organization_id')
    username = user_to_reactivate.get('username', 'N/A')

    # Varmista, että käyttäjä on inaktiivinen
    if user_to_reactivate.get('status') == 'active':
         flash(f"Käyttäjä {username} on jo aktiivinen.", "info")
         # Ohjaa silti takaisin, koska toiminto "onnistui" tavallaan
         return redirect(request.referrer or (url_for('superuser_view_organization_route', org_id=org_id) if org_id else url_for('superuser_orgs_route')))


    success, error_msg = db_manager.reactivate_user(user_id)

    if success:
        flash(f'Käyttäjä #{user_id} ({username}) aktivoitu onnistuneesti.', 'success')
        app.logger.info(f"Superuser {current_user.username} reactivated user {user_id} ('{username}')")
    else:
        flash(f'Virhe käyttäjän aktivoinnissa: {error_msg}', 'danger')
        app.logger.error(f"User reactivate error by superuser for ID {user_id}: {error_msg}")

    # Ohjaa takaisin käyttäjälistaan
    if org_id:
        return redirect(url_for('superuser_view_organization_route', org_id=org_id))
    else:
        # Jos käyttäjällä ei ollut organisaatiota
        return redirect(url_for('superuser_orgs_route'))


@app.route("/superuser/hard_delete_user/<int:user_id>", methods=['POST'])
@superuser_required
def superuser_hard_delete_user_route(user_id):
    """Pysyvä poisto (Hard Delete). VAIN SUPERUSER."""
    # Hae käyttäjä ensin varmistusta ja lokitusta varten
    user = db_manager.get_user_by_id(user_id)
    if not user:
        flash("Käyttäjää ei löytynyt.", "danger")
        return redirect(request.referrer or url_for('superuser_orgs_route'))

    org_id = user.get('organization_id')
    username = user.get('username', 'N/A')
    confirmation = request.form.get('confirmation')

    # Turvallisuusvarmistus: Varmista, että vahvistusteksti täsmää
    expected_confirmation = f"POISTA/{username}"
    if confirmation != expected_confirmation:
        flash(f"Vahvistus epäonnistui. Kirjoita '{expected_confirmation}' tarkalleen oikein.", "danger")
        return redirect(request.referrer or (url_for('superuser_view_organization_route', org_id=org_id) if org_id else url_for('superuser_orgs_route')))


    # Estä itsensä tai alkuperäisen superuserin (ID 1) poisto
    if user_id == current_user.id:
        flash("Et voi poistaa itseäsi pysyvästi.", "danger")
        return redirect(request.referrer or url_for('superuser_orgs_route'))
    # Lisää tarkistus ID 1:lle tai superuser-roolille, jos haluat estää sen poiston
    if user.get('role') == 'superuser': # Estä muiden superusereiden poisto
         flash("Toista pääkäyttäjää ei voi poistaa tätä kautta.", "danger")
         return redirect(request.referrer or url_for('superuser_orgs_route'))

    # Suorita pysyvä poisto
    app.logger.warning(f"Superuser {current_user.username} aloittaa KÄYTTÄJÄN {user_id} ('{username}') PYSYVÄN POISTON!")
    success, error_msg = db_manager.hard_delete_user(user_id)

    if success:
        flash(f'Käyttäjä #{user_id} ({username}) ja kaikki hänen tietonsa on POISTETTU PYSYVÄSTI.', 'warning')
        app.logger.critical(f"Superuser {current_user.username} HARD DELETED user {user_id} ('{username}') from org {org_id}") # Käytä CRITICAL-tasoa
    else:
        flash(f'KRIITTINEN VIRHE käyttäjän pysyvässä poistossa: {error_msg}', 'danger')
        app.logger.error(f"User hard delete FAILED for ID {user_id}: {error_msg}")

    # Ohjaa takaisin organisaationäkymään, jos mahdollista
    if org_id:
        return redirect(url_for('superuser_view_organization_route', org_id=org_id))
    else:
        # Jos käyttäjällä ei ollut organisaatiota
        return redirect(url_for('superuser_orgs_route'))


# --- Globaalien kysymysten hallintareitit (Siirretty tänne) ---

@app.route("/superuser/questions")
@superuser_required
def superuser_questions_route():
     """Näyttää kaikki kysymykset superuserille."""
     try:
          # Hae kaikki kysymykset, mukaan lukien 'needs_review'
          questions = db_manager._execute("SELECT * FROM questions ORDER BY status, category, id", fetch='all')
          return render_template("superuser_questions.html", questions=questions) # Tarvitaan uusi template
     except Exception as e:
          flash(f"Virhe kysymysten haussa: {e}","danger")
          return redirect(url_for('admin_route')) # Superuser dashboard

@app.route("/superuser/question/add", methods=['GET', 'POST'])
@superuser_required
def superuser_add_question_route():
    """Lomake ja logiikka uuden kysymyksen lisäämiseen."""
    if request.method == 'POST':
        question_text = request.form.get('question', '').strip()
        explanation = request.form.get('explanation', '').strip()
        # Logiikka uusien kategorioiden käsittelyyn
        category = request.form.get('new_category') if request.form.get('category') == '__add_new__' else request.form.get('category')
        difficulty = request.form.get('difficulty')
        options_list = [
            request.form.get(f'option_{i}', '').strip() for i in range(4)
        ]
        correct_answer_text = request.form.get('correct_answer', '').strip()
        status = request.form.get('status', 'validated') # Superuser voi asettaa statuksen

        # Validoinnit
        if not all([question_text, explanation, category, difficulty, correct_answer_text]) or not all(options_list):
            flash('Kaikki kentät (kysymys, selitys, kategoria, vaikeus, kaikki 4 vaihtoehtoa, oikea vastaus) ovat pakollisia.', 'danger')
            # Palauta lomake täytetyillä tiedoilla
            return render_template("superuser_add_question.html", categories=db_manager.get_categories(), question=request.form), 400

        if correct_answer_text not in options_list:
            flash('Annettu oikea vastaus ei löydy vaihtoehdoista.', 'danger')
            return render_template("superuser_add_question.html", categories=db_manager.get_categories(), question=request.form), 400

        # Tarkista duplikaatti
        is_duplicate, existing = db_manager.check_question_duplicate(question_text)
        if is_duplicate:
            flash(f'Varoitus: Samanlainen kysymys (ID: {existing["id"]}, Kategoria: {existing["category"]}) on jo olemassa. Tarkista ennen tallennusta.', 'warning')
            # Voit silti sallia tallennuksen tai vaatia vahvistuksen

        # Määritä oikea indeksi
        try:
             # Sekoita vaihtoehdot TAI pidä ne annetussa järjestyksessä
             # Jos pidetään järjestyksessä:
             correct_index = options_list.index(correct_answer_text)
             final_options = options_list
        except ValueError:
             # Ei pitäisi tapahtua yllä olevan tarkistuksen vuoksi
             flash('Virhe oikean vastauksen indeksin määrittelyssä.', 'danger')
             return render_template("superuser_add_question.html", categories=db_manager.get_categories(), question=request.form), 500

        # Lisää kysymys tietokantaan (tarvitaan uusi metodi db_manageriin)
        # success, error_msg = db_manager.add_question(
        #      question=question_text, options=final_options, correct=correct_index,
        #      explanation=explanation, category=category, difficulty=difficulty, status=status
        # )

        # Simuloitu lisäys (korvaa oikealla kutsulla)
        success, error_msg = True, None
        print("Tässä lisättäisiin kysymys:", question_text, final_options, correct_index, category, difficulty, status)

        if success:
            flash('Uusi kysymys lisätty onnistuneesti!', 'success')
            app.logger.info(f"Superuser {current_user.username} added new question in category {category}")
            return redirect(url_for('superuser_questions_route'))
        else:
            flash(f'Virhe kysymyksen lisäämisessä: {error_msg}', 'danger')
            app.logger.error(f"Superuser question add failed: {error_msg}")
            # Palauta lomake täytetyillä tiedoilla
            return render_template("superuser_add_question.html", categories=db_manager.get_categories(), question=request.form), 500

    # GET-pyyntö: Näytä tyhjä lomake
    categories = db_manager.get_categories()
    return render_template("superuser_add_question.html", categories=categories) # Tarvitaan uusi template


@app.route("/superuser/question/edit/<int:question_id>", methods=['GET', 'POST'])
@superuser_required
def superuser_edit_question_route(question_id):
    """Lomake ja logiikka kysymyksen muokkaamiseen."""
    # Hae kysymys ensin
    question = db_manager.get_single_question_for_edit(question_id)
    if not question:
        flash("Kysymystä ei löytynyt.", "danger")
        return redirect(url_for('superuser_questions_route'))

    if request.method == 'POST':
        # Kerää tiedot lomakkeelta
        data = {
            'question': request.form.get('question','').strip(),
            'explanation': request.form.get('explanation','').strip(),
            'options': [request.form.get(f'option_{i}', '').strip() for i in range(len(question.get('options',[])))], # Ota vanha määrä
            'correct': request.form.get('correct'), # Tulee indeksinä 0, 1, 2, 3
            'category': request.form.get('new_category') if request.form.get('category') == '__add_new__' else request.form.get('category'),
            'difficulty': request.form.get('difficulty'),
            'status': request.form.get('status', 'needs_review') # Salli statuksen muokkaus
        }

        # Validoinnit
        try:
            data['correct'] = int(data['correct'])
            if not (0 <= data['correct'] < len(data['options'])): raise ValueError
        except (ValueError, TypeError):
             flash("Virheellinen oikean vastauksen indeksi.", "danger")
             return render_template("superuser_edit_question.html", question=question, categories=db_manager.get_categories()), 400

        if not all(data.values()) or not all(data['options']) or data['category'] is None or data['difficulty'] is None:
            flash('Kaikki kentät ovat pakollisia.', 'danger')
            # Palauta lomake täytetyillä tiedoilla (käytä data-dictiä)
            question.update(data) # Päivitä question-dict uusilla arvoilla
            return render_template("superuser_edit_question.html", question=question, categories=db_manager.get_categories()), 400

        # Päivitä kysymys tietokantaan
        success, error_msg = db_manager.update_question(question_id, data)

        if success:
            flash('Kysymys päivitetty onnistuneesti!', 'success')
            app.logger.info(f"Superuser {current_user.username} edited question {question_id}")
            return redirect(url_for('superuser_questions_route'))
        else:
            flash(f'Virhe kysymyksen päivityksessä: {error_msg}', 'danger')
            app.logger.error(f"Superuser question update failed for ID {question_id}: {error_msg}")
            # Palauta lomake täytetyillä tiedoilla
            question.update(data)
            return render_template("superuser_edit_question.html", question=question, categories=db_manager.get_categories()), 500

    # GET-pyyntö: Näytä lomake esitäytettynä
    categories = db_manager.get_categories()
    return render_template("superuser_edit_question.html", question=question, categories=categories) # Tarvitaan uusi template


@app.route("/superuser/question/delete/<int:question_id>", methods=['POST'])
@superuser_required
def superuser_delete_question_route(question_id):
    """Poistaa kysymyksen ja siihen liittyvät tiedot."""
    # Lisää vahvistusmekanismi (esim. checkbox tai tekstikenttä)
    confirmation = request.form.get('confirmation')
    question_to_delete = db_manager.get_single_question_for_edit(question_id) # Hae tietoja varten
    if not question_to_delete:
         flash("Kysymystä ei löytynyt.", "danger")
         return redirect(url_for('superuser_questions_route'))

    # Varmista vahvistus
    if confirmation != f"POISTA/{question_id}":
        flash("Vahvistus epäonnistui. Kirjoita 'POISTA/<kysymysID>' tarkalleen oikein.", "danger")
        return redirect(request.referrer or url_for('superuser_questions_route'))


    success, error_msg = db_manager.delete_question(question_id)

    if success:
        flash(f'Kysymys #{question_id} ("{question_to_delete["question"][:30]}...") poistettu onnistuneesti.', 'success')
        app.logger.warning(f"Superuser {current_user.username} deleted question {question_id}")
    else:
        flash(f'Virhe kysymyksen poistossa: {error_msg}', 'danger')
        app.logger.error(f"Superuser question delete failed for ID {question_id}: {error_msg}")

    return redirect(url_for('superuser_questions_route'))


# Siirrä myös bulk upload, find duplicates, clear database tänne @superuser_required alle
# Esimerkiksi:
@app.route("/superuser/bulk_upload", methods=['POST'])
@superuser_required
def superuser_bulk_upload_route():
    # Sama logiikka kuin vanhassa admin_bulk_upload_route, mutta kutsutaan superuserina
    # ... toteutus ...
    pass

@app.route("/superuser/find_duplicates", methods=['GET', 'POST'])
@superuser_required
def superuser_find_duplicates_route():
     # Sama logiikka kuin vanhassa admin_find_duplicates_route
     # ... toteutus ...
     pass

@app.route("/superuser/clear_database", methods=['POST'])
@superuser_required
@limiter.limit("1 per hour") # Erittäin vaarallinen, rajoita tiukasti
def superuser_clear_database_route():
     # Sama logiikka kuin vanhassa admin_clear_database_route
     # Varmista vahvistus ERITTÄIN huolellisesti
     confirmation = request.form.get('confirmation')
     if confirmation == "TYHJENNA KAIKKI KYSYMYKSET NYT":
          success, result = db_manager.clear_all_questions()
          if success:
               flash(f"KAIKKI {result.get('deleted_count',0)} KYSYMYSTÄ POISTETTU!", "danger")
               app.logger.critical(f"SUPERUSER {current_user.username} TYHJENSI KOKO KYSYMYSTIETOKANNAN!")
          else:
               flash(f"Tietokannan tyhjennys epäonnistui: {result}", "danger")
     else:
          flash("Vahvistus epäonnistui. Tietokantaa EI tyhjennetty.", "warning")

     return redirect(url_for('admin_route')) # Superuser dashboard

# ==============================================================================
# --- PDF & WORD EXPORT REITIT (VAIN Superuser?) ---
# ==============================================================================
# Nämä vaativat ReportLab ja python-docx kirjastot
# Päätä, ovatko nämä vain superuserin käytössä vai myös organisaation adminien

# Yksinkertainen export kaikille kysymyksille (Superuser)
@app.route("/superuser/export/pdf")
@superuser_required
def superuser_export_pdf_quick():
    if not REPORTLAB_AVAILABLE:
        flash("PDF-kirjastoa (ReportLab) ei ole asennettu.", "danger")
        return redirect(url_for('admin_route'))
    try:
        # Hae KAIKKI kysymykset
        questions_raw = db_manager._execute("SELECT * FROM questions ORDER BY category, id", fetch='all')
        if not questions_raw:
             flash("Ei kysymyksiä vietäväksi.", "warning"); return redirect(url_for('admin_route'))

        questions_list = []
        for q in questions_raw:
             q_dict = dict(q)
             try: q_dict['options'] = json.loads(q_dict.get('options','[]'))
             except: q_dict['options'] = ["Virhe", "vaihtoehdoissa"]
             questions_list.append(q_dict)

        # Luo PDF (käytä aiemmin luotua create_pdf_document funktiota)
        pdf_buffer = create_pdf_document(questions_list, include_answers=True) # Olettaa funktion olemassaolon

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f'LOVe_Kysymykset_Kaikki_{timestamp}.pdf'
        app.logger.info(f"Superuser {current_user.username} exported {len(questions_list)} questions to PDF")

        from flask import make_response
        response = make_response(pdf_buffer.getvalue())
        response.headers['Content-Type'] = 'application/pdf'
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"' # Lainausmerkit tiedostonimeen
        return response

    # app.py - JATKOA EDELLISEEN

    except Exception as e:
        flash(f'Virhe PDF-viennissä: {e}', 'danger')
        app.logger.error(f"Superuser PDF export error: {e}", exc_info=True)
        # Lisää traceback debug-tilassa
        if DEBUG_MODE:
            import traceback
            traceback.print_exc()
        return redirect(url_for('admin_route')) # Ohjaa superuser dashboardiin


# Voit lisätä vastaavat reitit Wordille (/superuser/export/word) ja JSONille (/superuser/export/json)
# käyttäen create_word_document -funktiota ja json.dumpsia.

# --- PDF & Word Apufunktiot (Tarvitaan, jos export-reitit ovat käytössä) ---
# Nämä vaativat ReportLab ja python-docx

def create_pdf_document(questions_list, include_answers=True, title="LOVe Enhanced Kysymyspankki"):
    """Luo PDF-dokumentin kysymyksistä (vaatii ReportLab)."""
    if not REPORTLAB_AVAILABLE:
        raise RuntimeError("ReportLab-kirjastoa ei ole asennettu.")

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=0.75*inch, bottomMargin=0.75*inch)
    styles = getSampleStyleSheet()
    # Lisää omia tyylejä tarvittaessa
    story = []

    # Otsikko ja tiedot
    story.append(Paragraph(title, styles['h1']))
    story.append(Paragraph(f"Luotu: {datetime.now().strftime('%d.%m.%Y %H:%M')}", styles['Normal']))
    story.append(Paragraph(f"Kysymyksiä yhteensä: {len(questions_list)}", styles['Normal']))
    story.append(Spacer(1, 0.2*inch))

    # Kysymykset läpi
    for idx, q in enumerate(questions_list, 1):
        # Kysymys
        q_text = f"<b>{idx}. {q.get('question', 'N/A')}</b> (ID: {q.get('id', 'N/A')}, Kat: {q.get('category', 'N/A')}, Vaikeus: {q.get('difficulty', 'N/A')})"
        story.append(Paragraph(q_text, styles['Normal']))
        story.append(Spacer(1, 0.1*inch))

        # Vaihtoehdot
        options = q.get('options', [])
        correct_idx = q.get('correct', -1)
        letters = ['A', 'B', 'C', 'D']
        for i, option in enumerate(options):
            prefix = f"{letters[i]}. "
            style = styles['Normal']
            if include_answers and i == correct_idx:
                prefix = f"<b>{letters[i]}. ✓ </b>" # Merkitse oikea
                # Voit myös käyttää eri tyyliä oikealle vastaukselle
            option_para = Paragraph(f"{prefix}{option}", style)
            # Lisää sisennys
            option_para.style.leftIndent = 20
            story.append(option_para)

        # Selitys
        if include_answers and 'explanation' in q:
            story.append(Spacer(1, 0.05*inch))
            exp_text = f"<i>Selitys: {q['explanation']}</i>"
            exp_para = Paragraph(exp_text, styles['Italic'])
            exp_para.style.leftIndent = 10
            story.append(exp_para)

        story.append(Spacer(1, 0.15*inch))
        # Lisää sivunvaihto tarvittaessa
        # if idx % 5 == 0: story.append(PageBreak())

    try:
        doc.build(story)
        buffer.seek(0)
        return buffer
    except Exception as pdf_e:
        app.logger.error(f"Virhe PDF-dokumentin luonnissa: {pdf_e}", exc_info=True)
        # Voit heittää virheen tai palauttaa None/tyhjän bufferin
        raise RuntimeError(f"PDF-luonti epäonnistui: {pdf_e}")


def create_word_document(questions_list, include_answers=True, title="LOVe Enhanced Kysymyspankki"):
    """Luo Word-dokumentin kysymyksistä (vaatii python-docx)."""
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx-kirjastoa ei ole asennettu.")

    doc = Document()
    # Aseta marginaalit tms.
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Luotu: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    doc.add_paragraph(f"Kysymyksiä yhteensä: {len(questions_list)}")
    doc.add_paragraph()

    for idx, q in enumerate(questions_list, 1):
        # Kysymys
        p = doc.add_paragraph()
        p.add_run(f"{idx}. {q.get('question', 'N/A')}").bold = True
        p.add_run(f" (ID: {q.get('id', 'N/A')}, Kat: {q.get('category', 'N/A')}, Vaikeus: {q.get('difficulty', 'N/A')})").font.size = Pt(9)

        # Vaihtoehdot
        options = q.get('options', [])
        correct_idx = q.get('correct', -1)
        letters = ['A', 'B', 'C', 'D']
        for i, option in enumerate(options):
            opt_p = doc.add_paragraph(f"{letters[i]}. {option}", style='List Bullet 2') # Käytä sopivaa listatyyliä
            if include_answers and i == correct_idx:
                opt_p.runs[0].bold = True # Korosta oikea vastaus

        # Selitys
        if include_answers and 'explanation' in q:
            exp_p = doc.add_paragraph()
            exp_p.add_run("Selitys: ").italic = True
            exp_p.add_run(q['explanation']).italic = True
            exp_p.paragraph_format.left_indent = Inches(0.25)

        doc.add_paragraph() # Tyhjä rivi väliin

    buffer = BytesIO()
    try:
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    except Exception as docx_e:
        app.logger.error(f"Virhe Word-dokumentin luonnissa: {docx_e}", exc_info=True)
        raise RuntimeError(f"Word-luonti epäonnistui: {docx_e}")


# ==============================================================================
# --- VIRHEKÄSITTELY ---
# ==============================================================================
@app.errorhandler(404)
def not_found_error(error):
    """Käsittelee 404 Not Found -virheet."""
    # Lokita pyyntö, jotta voit tutkia virheellisiä linkkejä
    app.logger.warning(f"Sivua ei löytynyt (404): {request.url} (Viittaaja: {request.referrer})")
    # Näytä käyttäjäystävällinen 404-sivu
    return render_template('404.html'), 404

@app.errorhandler(500)
def internal_error(error):
    """Käsittelee 500 Internal Server Error -virheet."""
    # Yritä peruuttaa mahdollinen keskeneräinen tietokantatransaktio
    try:
        # Huom: Tämä ei välttämättä toimi kaikissa tilanteissa, riippuu virheen syystä
        # db_manager.rollback() # Tarvitset tällaisen metodin DatabaseManageriin
        pass # Tai jätä tietokannan hallintaan
    except Exception as rb_e:
        app.logger.error(f"Virhe rollbackin yrityksessä 500-käsittelijässä: {rb_e}")

    # Lokita virhe mahdollisimman tarkasti (sis. traceback)
    app.logger.error(f"Palvelinvirhe (500): {error} Polku: {request.url}", exc_info=True)

    # Näytä käyttäjäystävällinen 500-sivu
    # ÄLÄ näytä yksityiskohtaista virhettä käyttäjälle tuotannossa!
    return render_template('500.html'), 500

@app.errorhandler(403)
def forbidden_error(error):
    """Käsittelee 403 Forbidden -virheet."""
    user_info = f"Käyttäjä {getattr(current_user,'id','N/A')}" if current_user.is_authenticated else "Anonyymi käyttäjä"
    app.logger.warning(f"Pääsy estetty (403): {user_info} yritti polkuun {request.url}")
    # Näytä käyttäjäystävällinen 403-sivu
    return render_template('403.html'), 403

@app.errorhandler(429)
def ratelimit_handler(e):
    """Käsittelee rate limitin ylitykset (429 Too Many Requests)."""
    # Lokita tapahtuma
    app.logger.warning(f"Rate limit ylitetty: {request.remote_addr} polkuun {request.path}. Raja: {e.description}")

    # Jos pyyntö oli API:lle, palauta JSON
    # Tarkista 'Accept' header tai polun alku '/api/'
    if request.path.startswith('/api/') or (request.accept_mimetypes and request.accept_mimetypes.accept_json and not request.accept_mimetypes.accept_html):
        return jsonify(error="Liikaa pyyntöjä.", limit=e.description), 429

    # Muuten näytä HTML-sivu
    # Varmista, että templates/429.html on olemassa
    return render_template('429.html', limit=e.description), 429


# ==============================================================================
# --- SOVELLUKSEN KÄYNNISTYS ---
# ==============================================================================
if __name__ == '__main__':
    # Hae portti ympäristömuuttujasta (oletus 8080)
    # Tärkeää esim. Railway, Heroku yms. varten
    port = int(os.environ.get("PORT", 8080))

    # Käynnistä Flaskin oma kehityspalvelin VAIN JOS ajetaan suoraan JA DEBUG_MODE on päällä
    # Tuotannossa Gunicorn/uWSGI hoitaa tämän osion.
    if DEBUG_MODE:
        app.logger.info(f"Käynnistetään Flaskin kehityspalvelin osoitteessa http://0.0.0.0:{port}/ (DEBUG MODE)")
        # host='0.0.0.0' sallii yhteydet paikallisverkon muista koneista
        # threaded=True voi auttaa, jos käsittelet useita pyyntöjä samanaikaisesti kehityksessä
        app.run(host='0.0.0.0', port=port, debug=True, threaded=True)
    else:
        # Tuotantotilassa Gunicorn/uWSGI importtaa 'app'-olion tästä tiedostosta.
        # Tämä lohko ei yleensä suoritu silloin.
        app.logger.info("Sovellus on valmis ajettavaksi tuotantopalvelimella (Gunicorn/uWSGI).")
        # Voit lisätä tähän esim. tietokantayhteyden testauksen ennen kuin Gunicorn ottaa ohjat.
        try:
             # Testaa tietokantayhteys
             conn_test = db_manager.get_connection()
             conn_test.close()
             app.logger.info("Tietokantayhteys testattu onnistuneesti käynnistyksessä.")
        except Exception as db_test_e:
             app.logger.critical(f"KRIITTINEN: Tietokantayhteys EI TOIMI käynnistyksessä: {db_test_e}", exc_info=True)
             # Voit pysäyttää tässä, jotta Gunicorn ei käynnisty turhaan
             # import sys
             # sys.exit("Tietokantayhteys epäonnistui.")

# Tiedoston loppu

